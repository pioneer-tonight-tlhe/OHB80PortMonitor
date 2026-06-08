from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
from pathlib import Path

root = Path(r"D:\Project\CYTC_Project\OHB80PortMonitor\OHB80PortMonitor_V_1_0_0")
assets = root / "docs" / "manual_assets"
out_doc = root / "docs" / "OHB80PortMonitor_用户使用手册_v1.0.docx"

images = [
    assets / "img_01.png",
    assets / "img_02.png",
    assets / "img_03.png",
    assets / "img_04.png",
    assets / "img_05.png",
    assets / "img_06.png",
    assets / "img_07.png",
    assets / "img_08.png",
    assets / "img_09.png",
]

for p in images:
    if not p.exists():
        raise FileNotFoundError(f"missing image: {p}")

doc = Document()

section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
r = p.add_run("OHB80PortMonitor 用户使用手册")
r.bold = True
r.font.size = Pt(24)

p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
r = p.add_run("版本：V1.0.0")
r.font.size = Pt(14)

p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
r = p.add_run(f"编制日期：{datetime.now().strftime('%Y-%m-%d')}")
r.font.size = Pt(12)

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph("文档用途")
p.runs[0].bold = True
p.runs[0].font.size = Pt(14)

doc.add_paragraph("本手册用于指导操作人员完成 OHB80PortMonitor 软件的登录、监控查看、页面导航与配置操作。")
doc.add_paragraph("本文配图来自当前版本软件界面，可用于培训、交接与现场操作参考。")

doc.add_page_break()

fig_no = 1

def heading(text, level=1):
    doc.add_heading(text, level=level)

def add_figure(path, caption, width):
    global fig_no
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cp = doc.add_paragraph(f"图{fig_no}  {caption}")
    cp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    fig_no += 1

heading("1. 主界面总览", 1)
doc.add_paragraph("主界面由顶部状态栏、左侧导航栏、中部监控表格和下方轨道图组成。")
add_figure(images[0], "主界面总览（Home + Alarm）", 6.2)

heading("2. 账号登录与权限", 1)
doc.add_paragraph("系统支持多级权限控制。未登录时仅显示游客权限，登录后可根据账号等级开放配置/调试等功能。")
heading("2.1 未登录状态", 2)
add_figure(images[1], "账号菜单（未登录，Level: Guest）", 3.0)
heading("2.2 登录窗口", 2)
doc.add_paragraph("输入用户名与密码后，点击 Login 完成认证。")
add_figure(images[2], "用户登录窗口", 4.8)
heading("2.3 登录后状态", 2)
add_figure(images[3], "账号菜单（已登录，Level: Normal）", 3.0)

heading("3. 首页监控页面（Home）", 1)
doc.add_paragraph("首页用于实时查看设备运行参数，包括 QRCode、压力、流量、湿度、温度等。")
add_figure(images[4], "首页监控总览（含左侧导航）", 6.2)

heading("4. 视图切换操作", 1)
doc.add_paragraph("通过右上角 Foup View / Set View 按钮可切换轨道图展示方式。")
add_figure(images[5], "Set View 视图示例", 6.2)
add_figure(images[6], "Foup View 视图示例", 6.2)

heading("5. 配置页面（Config）", 1)
doc.add_paragraph("配置页按功能分区，支持 Idle Purge、Pneumatic Valve、SH85 Periodic、SH85 Self-check 等参数设置。")
add_figure(images[7], "配置页总览", 6.2)

heading("6. SH85 自检功能", 1)
doc.add_paragraph("SH85 自检包含两部分：定期自检（Periodic）和手动自检（Self-check）。")
doc.add_paragraph("1. 定期自检：设置周期并点击 Set，状态栏显示下一次检查倒计时。")
doc.add_paragraph("2. 手动自检：输入 Target Device ID，点击 Check，等待约 70 秒完成。")
add_figure(images[8], "SH85 自检区域特写（Periodic + Manual）", 6.0)

heading("7. 常见操作建议", 1)
doc.add_paragraph("1. 登录后先确认右上角账号状态与权限等级。")
doc.add_paragraph("2. 参数修改前先确认目标设备 ID/QRCode，避免误操作。")
doc.add_paragraph("3. 配置生效后观察顶部告警条与首页参数回读结果。")
doc.add_paragraph("4. 出现异常时优先检查网络配置、设备在线状态和权限范围。")

heading("8. 版本说明", 1)
doc.add_paragraph("手册版本：V1.0.0")
doc.add_paragraph("适用软件：OHB80PortMonitor V1.0.0")
doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")

doc.save(out_doc)
print(out_doc)

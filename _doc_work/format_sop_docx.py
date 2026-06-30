from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


SOURCE = Path(r"D:\Project\CYTC_Project\OHB80PortMonitor\_doc_work\OHB80PortMonitor_SOP_original.docx")
OUTPUT = Path(r"D:\Project\CYTC_Project\OHB80PortMonitor\_doc_work\OHB80PortMonitor_SOP_format_optimized.docx")

BODY_FONT = "Microsoft YaHei"
ACCENT = RGBColor(31, 77, 120)
HEADING = RGBColor(46, 116, 181)
MUTED = RGBColor(89, 89, 89)
TABLE_HEADER_FILL = "E8EEF5"
TABLE_BORDER = "B7C6D8"


def set_run_font(run, name=BODY_FONT, size=None, bold=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_style_font(style, size, bold=False, color=None):
    font = style.font
    font.name = BODY_FONT
    style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    style._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    font.size = Pt(size)
    font.bold = bold
    if color is not None:
        font.color.rgb = color


def get_or_create_paragraph_style(doc, name, base="Normal"):
    try:
        return doc.styles[name]
    except KeyError:
        style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles[base]
        return style


def set_outline_level(style, level):
    p_pr = style._element.get_or_add_pPr()
    outline = p_pr.find(qn("w:outlineLvl"))
    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        p_pr.append(outline)
    outline.set(qn("w:val"), str(level))


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.25):
    fmt = paragraph.paragraph_format
    fmt.left_indent = Pt(0)
    fmt.right_indent = Pt(0)
    fmt.first_line_indent = Pt(0)
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def clear_paragraph(paragraph):
    p = paragraph._element
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)


def add_field(paragraph, instr_text):
    run_begin = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run_begin._r.append(fld_char_begin)

    run_instr = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instr_text
    run_instr._r.append(instr)

    run_sep = paragraph.add_run()
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    run_sep._r.append(fld_char_sep)

    run_text = paragraph.add_run("1")
    set_run_font(run_text, size=9, color=MUTED)

    run_end = paragraph.add_run()
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run_end._r.append(fld_char_end)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            if index < len(widths_dxa):
                set_cell_width(cell, widths_dxa[index])


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "6")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), TABLE_BORDER)


def remove_empty_body_paragraphs(doc):
    for paragraph in list(doc.paragraphs):
        has_text = bool(paragraph.text.strip())
        has_drawing = bool(paragraph._p.xpath(".//w:drawing"))
        if not has_text and not has_drawing:
            paragraph._element.getparent().remove(paragraph._element)


def has_drawing(paragraph):
    return bool(paragraph._p.xpath(".//w:drawing"))


def is_complete_sentence(text):
    return text.endswith(("。", "：", ":", "；", ";", "！", "？", ")", "）"))


def should_merge_with_next(current, next_paragraph):
    current_text = current.text.strip()
    next_text = next_paragraph.text.strip()
    if not current_text or not next_text:
        return False
    if has_drawing(current) or has_drawing(next_paragraph):
        return False
    if normalize_heading_text(current_text) or normalize_heading_text(next_text):
        return False
    if current_text == "如下：" or next_text == "如下：":
        return False
    if is_complete_sentence(current_text):
        return False
    return True


def merge_manual_line_breaks(doc):
    paragraphs = list(doc.paragraphs)
    index = 0
    while index < len(paragraphs) - 1:
        current = paragraphs[index]
        next_paragraph = paragraphs[index + 1]
        if should_merge_with_next(current, next_paragraph):
            current.text = current.text.rstrip() + next_paragraph.text.lstrip()
            next_paragraph._element.getparent().remove(next_paragraph._element)
            paragraphs.pop(index + 1)
            continue
        index += 1


def apply_section_layout(doc):
    for section in doc.sections:
        section.start_type = WD_SECTION_START.NEW_PAGE
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.header_distance = Cm(1.0)
        section.footer_distance = Cm(0.9)

        header = section.header
        header.is_linked_to_previous = False
        paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        clear_paragraph(paragraph)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(paragraph, before=0, after=2, line=1.0)
        run = paragraph.add_run("OHB80PortMonitor 软件 SOP | Home 主界面")
        set_run_font(run, size=9, bold=False, color=MUTED)

        footer = section.footer
        footer.is_linked_to_previous = False
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        clear_paragraph(paragraph)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(paragraph, before=0, after=0, line=1.0)
        run = paragraph.add_run("第 ")
        set_run_font(run, size=9, color=MUTED)
        add_field(paragraph, "PAGE")
        run = paragraph.add_run(" 页")
        set_run_font(run, size=9, color=MUTED)


def apply_styles(doc):
    normal = doc.styles["Normal"]
    set_style_font(normal, 10.5, color=RGBColor(33, 33, 33))
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    h1 = get_or_create_paragraph_style(doc, "SOP Heading 1")
    set_style_font(h1, 17, bold=True, color=ACCENT)
    h1.paragraph_format.space_before = Pt(14)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.keep_with_next = True
    set_outline_level(h1, 0)

    h2 = get_or_create_paragraph_style(doc, "SOP Heading 2")
    set_style_font(h2, 14, bold=True, color=HEADING)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(7)
    h2.paragraph_format.keep_with_next = True
    set_outline_level(h2, 1)

    h3 = get_or_create_paragraph_style(doc, "SOP Heading 3")
    set_style_font(h3, 12, bold=True, color=ACCENT)
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(5)
    h3.paragraph_format.keep_with_next = True
    set_outline_level(h3, 2)

    caption = get_or_create_paragraph_style(doc, "SOP Caption")
    set_style_font(caption, 9, bold=False, color=MUTED)
    caption.paragraph_format.left_indent = Pt(0)
    caption.paragraph_format.right_indent = Pt(0)
    caption.paragraph_format.first_line_indent = Pt(0)
    caption.paragraph_format.space_before = Pt(2)
    caption.paragraph_format.space_after = Pt(4)
    caption.paragraph_format.line_spacing = 1.15
    caption.paragraph_format.keep_with_next = True


def normalize_heading_text(text):
    mapping = {
        "4  软件介绍": ("4 软件介绍", "SOP Heading 1"),
        "4.1Home界面介绍": ("4.1 Home 界面介绍", "SOP Heading 2"),
        "4.1 登录/登出": ("4.1.1 登录 / 登出", "SOP Heading 3"),
        "4.2 运行日志公告栏": ("4.1.2 运行日志公告栏", "SOP Heading 3"),
        "4.3 Set详细信息": ("4.1.3 Set 详细信息", "SOP Heading 3"),
        "4.4 界面导航栏": ("4.1.4 界面导航栏", "SOP Heading 3"),
        "4.5 设备背景颜色": ("4.1.5 设备背景颜色", "SOP Heading 3"),
        "4.6 视图切换": ("4.1.6 视图切换", "SOP Heading 3"),
        "4.7 视图大小切换": ("4.1.7 视图大小切换", "SOP Heading 3"),
    }
    return mapping.get(text.strip())


def format_body(doc):
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        heading = normalize_heading_text(text)
        if heading:
            new_text, style = heading
            paragraph.text = new_text
            paragraph.style = style
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            continue

        if "在4.6具体解释" in paragraph.text:
            paragraph.text = paragraph.text.replace("在4.6具体解释", "在 4.1.6 中具体解释")

        if has_drawing(paragraph):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(paragraph, before=4, after=8, line=1.0)
        elif text == "如下：":
            paragraph.style = "SOP Caption"
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.keep_with_next = True
            set_paragraph_spacing(paragraph, before=2, after=4, line=1.15)
            for run in paragraph.runs:
                set_run_font(run, size=9, color=MUTED)
        else:
            paragraph.style = "Normal"
            set_paragraph_spacing(paragraph, before=0, after=6, line=1.25)
            for run in paragraph.runs:
                set_run_font(run, size=10.5, color=RGBColor(33, 33, 33))


def format_images(doc):
    max_width = Inches(6.15)
    for shape in doc.inline_shapes:
        if shape.width > max_width:
            ratio = max_width / shape.width
            shape.width = max_width
            shape.height = int(shape.height * ratio)


def format_tables(doc):
    width_patterns = {
        2: [2450, 6790],
        3: [2100, 3200, 3940],
    }
    for table in doc.tables:
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.autofit = False
        set_table_borders(table)
        widths = width_patterns.get(len(table.columns), [int(9240 / len(table.columns))] * len(table.columns))
        set_table_geometry(table, widths)

        for row_index, row in enumerate(table.rows):
            for col_index, cell in enumerate(row.cells):
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                set_cell_margins(cell)
                if row_index == 0:
                    shade_cell(cell, TABLE_HEADER_FILL)
                for paragraph in cell.paragraphs:
                    set_paragraph_spacing(paragraph, before=0, after=0, line=1.2)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 or col_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        set_run_font(run, size=9.5, bold=(row_index == 0), color=RGBColor(33, 33, 33))


def main():
    doc = Document(SOURCE)
    apply_section_layout(doc)
    apply_styles(doc)
    merge_manual_line_breaks(doc)
    remove_empty_body_paragraphs(doc)
    format_body(doc)
    format_images(doc)
    format_tables(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()

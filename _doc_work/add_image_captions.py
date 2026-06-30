from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph


SOURCE = Path(r"D:\Project\CYTC_Project\OHB80PortMonitor\_doc_work\OHB80PortMonitor_SOP_home_synced.docx")
OUTPUT = Path(r"D:\Project\CYTC_Project\OHB80PortMonitor\_doc_work\OHB80PortMonitor_SOP_image_captions.docx")


CAPTIONS = [
    "图 4-1 Home 主界面整体布局与编号标注（正文中的【1】至【9】均对应本图编号）",
    "图 4-2 未登录状态下的用户账号信息弹框",
    "图 4-3 用户登录界面",
    "图 4-4 登录后的用户账号信息弹框",
    "图 4-5 运行日志实时记录界面（Live Log）",
    "图 4-6 运行日志历史记录查询界面（History）",
    "图 4-7 Set 视图下的 OHB 设备实时数据表",
    "图 4-8 登录后的界面导航栏权限显示效果",
    "图 4-9 Foup View 模式下的视图区域效果",
    "图 4-10 视图区域放大效果",
    "图 4-11 视图区域缩小效果",
    "图 4-12 Cabinet 界面整体布局与编号标注（正文中的【1】至【6】均对应本图编号）",
]


TEXT_REPLACEMENTS = {
    "运行日志公告栏中会实时滚动显示最新的运行日汇总信息。": "运行日志公告栏中会实时滚动显示最新的运行日志汇总信息。",
    "根据图中编号，Home 界面主要区域说明如下：": "根据图 4-1 中的编号，Home 界面主要区域说明如下：",
    "点击【1】用户头像，弹出用户账号信息弹框，": "点击图 4-1 中【1】用户头像，弹出用户账号信息弹框，",
    "输入正确的账号密码，点击Login完成登录，登录后再次点击【1】用户头像，用户账号信息弹框，": "输入正确的账号密码，点击Login完成登录；登录后再次点击图 4-1 中【1】用户头像，会显示用户账号信息弹框，",
    "点击【2】运行日志公告栏，弹出运行日志界面（默认打开实时界面），": "点击图 4-1 中【2】运行日志公告栏，弹出运行日志界面（默认打开实时界面），",
    "未登录账号，只显示图片【4】中3个界面：Home（主界面）、Cabinet（电控柜界面）、Alarm（警报界面）。": "未登录账号时，只显示图 4-1 中【4】导航栏内的 3 个界面：Home（主界面）、Cabinet（电控柜界面）、Alarm（警报界面）。",
    "【5】中的设备背景颜色代表着设备当前所处的不同的状态，颜色状态表，": "图 4-1 中【5】的设备背景颜色代表设备当前所处的不同状态，颜色状态表，",
    "当点击【6】中的Foup View按钮，【8】视图区域会转换成Foup视图模式模式，每1个Set控件会被展开成4个Foup控件,界面效果，": "当点击图 4-1 中【6】的 Foup View 按钮，图 4-1 中【8】视图区域会转换成 Foup 视图模式，每 1 个 Set 控件会被展开成 4 个 Foup 控件，界面效果，",
    "当点击【7】的【+】或者【-】按钮，能够控制【8】视图区域的伸缩，界面放大效果，如下：": "当点击图 4-1 中【7】的【+】或者【-】按钮，能够控制图 4-1 中【8】视图区域的伸缩，界面放大效果如下：",
    "根据图中编号，Cabinet 界面主要区域说明如下：": "根据图 4-12 中的编号，Cabinet 界面主要区域说明如下：",
}


def has_image(paragraph: Paragraph) -> bool:
    return bool(paragraph._element.xpath(".//w:drawing | .//w:pict"))


def insert_paragraph_after(paragraph: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def set_caption_font(run) -> None:
    run.font.name = "Microsoft YaHei"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(90, 90, 90)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run._element.rPr.rFonts.set(qn("w:ascii"), "Microsoft YaHei")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Microsoft YaHei")


def replace_paragraph_text(paragraph: Paragraph, new_text: str) -> None:
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = new_text
    else:
        paragraph.add_run(new_text)


def apply_caption_style(paragraph: Paragraph) -> None:
    try:
        paragraph.style = "Caption"
    except Exception:
        pass
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.0


def main() -> None:
    doc = Document(SOURCE)

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text in TEXT_REPLACEMENTS:
            replace_paragraph_text(paragraph, TEXT_REPLACEMENTS[text])

    image_paragraphs = [paragraph for paragraph in doc.paragraphs if has_image(paragraph)]
    if len(image_paragraphs) != len(CAPTIONS):
        raise RuntimeError(f"Expected {len(CAPTIONS)} images, found {len(image_paragraphs)}")

    for paragraph, caption_text in reversed(list(zip(image_paragraphs, CAPTIONS))):
        paragraph.paragraph_format.keep_with_next = True
        caption = insert_paragraph_after(paragraph)
        apply_caption_style(caption)
        run = caption.add_run(caption_text)
        set_caption_font(run)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()

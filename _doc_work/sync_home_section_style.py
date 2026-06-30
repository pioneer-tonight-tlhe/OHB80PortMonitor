from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


SOURCE = Path(r"D:\Project\CYTC_Project\OHB80PortMonitor\_doc_work\OHB80PortMonitor_SOP_with_cabinet.docx")
OUTPUT = Path(r"D:\Project\CYTC_Project\OHB80PortMonitor\_doc_work\OHB80PortMonitor_SOP_home_synced.docx")

BODY_FONT = "Microsoft YaHei"
ACCENT = RGBColor(31, 77, 120)
HEADING = RGBColor(46, 116, 181)
MUTED = RGBColor(89, 89, 89)
INK = RGBColor(33, 33, 33)
TABLE_HEADER_FILL = "E8EEF5"
TABLE_BORDER = "B7C6D8"


def set_run_font(run, name=BODY_FONT, size=None, bold=None, color=None):
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:eastAsia"), name)
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.25):
    fmt = paragraph.paragraph_format
    fmt.left_indent = Pt(0)
    fmt.right_indent = Pt(0)
    fmt.first_line_indent = Pt(0)
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def style_paragraph(paragraph, before=0, after=6, line=1.25, color=INK, size=10.5, bold=False):
    set_paragraph_spacing(paragraph, before=before, after=after, line=line)
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=bold, color=color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=140, bottom=90, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
    cant_split.set(qn("w:val"), "true")


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "6")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), TABLE_BORDER)


def set_table_geometry(table, widths_dxa):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            if index < len(widths_dxa):
                set_cell_width(cell, widths_dxa[index])


def format_table(table, widths_dxa):
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.autofit = False
    set_table_borders(table)
    set_table_geometry(table, widths_dxa)
    if table.rows:
        set_repeat_header(table.rows[0])
    for row_index, row in enumerate(table.rows):
        set_cant_split(row)
        for col_index, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            if row_index == 0:
                shade_cell(cell, TABLE_HEADER_FILL)
            for paragraph in cell.paragraphs:
                set_paragraph_spacing(paragraph, before=0, after=0, line=1.18)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 or col_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    set_run_font(run, size=9.3 if row_index else 9.5, bold=(row_index == 0), color=INK)


def insert_after(reference_element, new_element):
    parent = reference_element.getparent()
    parent.insert(parent.index(reference_element) + 1, new_element)
    return new_element


def move_after(reference_element, moving_element):
    parent = moving_element.getparent()
    parent.remove(moving_element)
    return insert_after(reference_element, moving_element)


def make_paragraph(doc, text, style=None, before=0, after=6, line=1.25, keep_next=False, color=INK, size=10.5):
    paragraph = doc.add_paragraph()
    if style is not None:
        paragraph.style = style
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=color)
    set_paragraph_spacing(paragraph, before=before, after=after, line=line)
    paragraph.paragraph_format.keep_with_next = keep_next
    return paragraph


def make_heading(doc, text, style):
    paragraph = doc.add_paragraph(text)
    paragraph.style = style
    paragraph.paragraph_format.keep_with_next = True
    for run in paragraph.runs:
        if style.endswith("2"):
            set_run_font(run, size=14, bold=True, color=HEADING)
        else:
            set_run_font(run, size=12, bold=True, color=ACCENT)
    return paragraph


def make_home_area_table(doc):
    rows = [
        ("1", "登录 / 登出", "显示当前用户账号信息、登录入口和登出入口。未登录时仅允许查看基础界面，登录后根据权限开放更多功能。"),
        ("2", "运行日志公告栏", "实时滚动显示最新运行日志汇总信息；点击后可打开运行日志窗口，查看实时日志或历史记录。"),
        ("3", "Set 详细信息", "显示当前 Set 视图下选中区域的 OHB 设备实时数据，例如二维码、进出气压力、流量、湿度和温度等。"),
        ("4", "界面导航栏", "用于切换 Home、Cabinet、Alarm 等界面；登录后按权限显示 Config、Comm. 等更多入口。"),
        ("5", "设备背景颜色", "通过不同背景颜色标识 OHB 设备当前状态，例如充氮时间、Foup Out、Alarm 和 Disable 等状态。"),
        ("6", "视图切换", "用于在 Foup View 与 Set View 之间切换，使界面按 Foup 或 Set 维度展示设备状态。"),
        ("7", "视图大小设置", "通过【+】和【-】调整视图区域缩放比例，便于查看更多设备或放大局部状态。"),
        ("8", "视图区域", "展示 OHB 设备布局和设备实时状态，是 Home 界面主要监控区域。"),
        ("9", "Set 视图设备", "在 Set View 模式下显示 Set 相关设备卡片，可用于快速查看设备编号、状态和关键实时数据。"),
    ]
    table = doc.add_table(rows=1, cols=3)
    for index, text in enumerate(("编号", "区域", "功能说明")):
        table.cell(0, index).text = text
    for row in rows:
        cells = table.add_row().cells
        for index, text in enumerate(row):
            cells[index].text = text
    format_table(table, [780, 2650, 5930])
    return table


def text_of_element(element):
    return "".join(t.text or "" for t in element.xpath(".//w:t")).strip()


def sync_home_section(doc):
    body = doc.element.body
    children = list(body.iterchildren())
    home_heading = None
    home_screenshot = None

    for child in children:
        if child.tag == qn("w:p"):
            text = text_of_element(child)
            if text == "4.1 Home 界面介绍":
                home_heading = child
            elif home_screenshot is None and child.xpath(".//w:drawing"):
                # The first screenshot is the annotated Home overview image.
                home_screenshot = child

    if home_heading is None or home_screenshot is None:
        raise RuntimeError("Could not locate Home heading or overview screenshot.")

    # Make the general software introduction independent from the moved screenshot.
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith("双击OHB80PortMonitor.exe"):
            paragraph.text = "双击 OHB80PortMonitor.exe 打开软件，与机台正常连接后会进入 Home 主界面。未登录前，无法执行任何控制操作。"
            style_paragraph(paragraph, after=8)
            break

    ref = home_heading
    overview = make_paragraph(
        doc,
        "概述：Home 界面用于查看 OHB 设备实时状态、运行日志公告、界面导航、设备状态颜色说明，并提供 Set / Foup 视图切换和视图缩放入口。",
        after=8,
    )
    ref = insert_after(ref, overview._element)

    overview_caption = make_paragraph(doc, "Home 界面整体如下：", after=4, keep_next=True)
    ref = insert_after(ref, overview_caption._element)

    ref = move_after(ref, home_screenshot)

    area_heading = make_heading(doc, "4.1.1 界面区域说明", "SOP Heading 3")
    ref = insert_after(ref, area_heading._element)

    area_intro = make_paragraph(doc, "根据图中编号，Home 界面主要区域说明如下：", after=4, keep_next=True)
    ref = insert_after(ref, area_intro._element)

    table = make_home_area_table(doc)
    ref = insert_after(ref, table._element)

    spacer = doc.add_paragraph()
    ref = insert_after(ref, spacer._element)

    replacements = {
        "4.1.1 登录 / 登出": "4.1.2 登录 / 登出",
        "4.1.2 运行日志公告栏": "4.1.3 运行日志公告栏",
        "4.1.3 Set 详细信息": "4.1.4 Set 详细信息",
        "4.1.4 界面导航栏": "4.1.5 界面导航栏",
        "4.1.5 设备背景颜色": "4.1.6 设备背景颜色",
        "4.1.6 视图切换": "4.1.7 视图切换",
        "4.1.7 视图大小切换": "4.1.8 视图大小切换",
    }

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text in replacements:
            paragraph.text = replacements[text]
            paragraph.style = "SOP Heading 3"
            style_paragraph(paragraph, color=ACCENT, size=12, bold=True)
        elif "在 4.1.6 中具体解释" in paragraph.text:
            paragraph.text = paragraph.text.replace("在 4.1.6 中具体解释", "在 4.1.7 中具体解释")
            style_paragraph(paragraph)


def main():
    doc = Document(SOURCE)
    sync_home_section(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()

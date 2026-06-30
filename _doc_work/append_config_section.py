# -*- coding: utf-8 -*-
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SOURCE = Path(r"D:\Project\CYTC_Project\OHB80PortMonitor\_doc_work\OHB80PortMonitor_SOP_image_captions.docx")
OUTPUT = Path(r"D:\Project\CYTC_Project\OHB80PortMonitor\_doc_work\OHB80PortMonitor_SOP_add_config.docx")

IMG_CONFIG_OVERVIEW = Path(r"C:\Users\Lenovo\AppData\Local\Temp\codex-clipboard-1b2b6ce4-64f4-43a1-82d1-e6dfc3bdcabb.png")
IMG_IDLE_PURGE = Path(r"C:\Users\Lenovo\AppData\Local\Temp\codex-clipboard-82de1339-05b8-4fc0-b799-fc3595a44c8c.png")
IMG_PNEUMATIC = Path(r"C:\Users\Lenovo\AppData\Local\Temp\codex-clipboard-2fe45d92-786a-44b1-aca3-990a50aa01e7.png")
IMG_PERIODIC = Path(r"C:\Users\Lenovo\AppData\Local\Temp\codex-clipboard-f9ae21a3-7754-4940-bdff-dbba4f908acb.png")
IMG_FLOW = Path(r"C:\Users\Lenovo\AppData\Local\Temp\codex-clipboard-f51f1636-9495-4b34-90cd-acf91a4d21d8.png")
IMG_REPORT_LIVE = Path(r"C:\Users\Lenovo\AppData\Local\Temp\codex-clipboard-5c031937-51d6-4e82-b02d-a55a9de28542.png")
IMG_REPORT_HISTORY = Path(r"C:\Users\Lenovo\AppData\Local\Temp\codex-clipboard-a7e60a1a-9979-49b5-a02a-caf17d79b1f6.png")
IMG_MANUAL_SELF_CHECK = Path(r"C:\Users\Lenovo\AppData\Local\Temp\codex-clipboard-e6ac95c5-1675-4d25-ac3e-19a1141c29ca.png")
IMG_PURGE_FLOW = Path(r"C:\Users\Lenovo\AppData\Local\Temp\codex-clipboard-fc4327de-4304-4ff8-8957-f8514c69a00c.png")
IMG_DEVICE_ENABLE = Path(r"C:\Users\Lenovo\AppData\Local\Temp\codex-clipboard-105c77cc-efe6-4325-ac1d-a3d7fb9d9ca0.png")

BODY_FONT = "Microsoft YaHei"
ACCENT = RGBColor(31, 77, 120)
HEADING = RGBColor(46, 116, 181)
MUTED = RGBColor(89, 89, 89)
INK = RGBColor(33, 33, 33)
TABLE_HEADER_FILL = "E8EEF5"
TABLE_BORDER = "B7C6D8"
CALLOUT_FILL = "F4F6F9"
STATUS_FILL = "FFF2CC"


def set_run_font(run, name=BODY_FONT, size=None, bold=None, color=None):
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:eastAsia"), name)
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.25):
    fmt = paragraph.paragraph_format
    fmt.left_indent = Pt(0)
    fmt.right_indent = Pt(0)
    fmt.first_line_indent = Pt(0)
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=140, bottom=90, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "6")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), TABLE_BORDER)


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def set_table_geometry(table, widths_dxa):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            if index < len(widths_dxa):
                set_cell_width(cell, widths_dxa[index])


def add_heading(doc, text, style_name, before=None):
    paragraph = doc.add_paragraph(text)
    paragraph.style = style_name
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.keep_with_next = True
    if before is not None:
        paragraph.paragraph_format.space_before = Pt(before)
    for run in paragraph.runs:
        if style_name.endswith("2"):
            set_run_font(run, size=14, bold=True, color=HEADING)
        else:
            set_run_font(run, size=12, bold=True, color=ACCENT)
    return paragraph


def add_body_paragraph(doc, text, before=0, after=6, keep_next=False):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5, color=INK)
    set_paragraph_spacing(paragraph, before=before, after=after, line=1.25)
    paragraph.paragraph_format.keep_with_next = keep_next
    return paragraph


def add_image(doc, image_path, width_inches=6.15):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    set_paragraph_spacing(paragraph, before=2, after=2, line=1.0)
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    return paragraph


def add_caption(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(paragraph, before=2, after=8, line=1.0)
    run = paragraph.add_run(text)
    set_run_font(run, size=9, color=MUTED)
    return paragraph


def format_table(table, widths_dxa, header_fill=TABLE_HEADER_FILL, font_size=9.2):
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.autofit = False
    set_table_borders(table)
    set_table_geometry(table, widths_dxa)
    if table.rows:
        set_repeat_header(table.rows[0])
    for row_index, row in enumerate(table.rows):
        set_cant_split(row)
        for col_index, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            if row_index == 0:
                shade_cell(cell, header_fill)
            for paragraph in cell.paragraphs:
                set_paragraph_spacing(paragraph, before=0, after=0, line=1.18)
                if row_index == 0 or col_index == 0:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    set_run_font(run, size=font_size if row_index else 9.5, bold=(row_index == 0), color=INK)


def add_table(doc, headers, rows, widths_dxa, header_fill=TABLE_HEADER_FILL, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    for index, text in enumerate(headers):
        table.cell(0, index).text = text
    for row in rows:
        cells = table.add_row().cells
        for index, text in enumerate(row):
            cells[index].text = text
    format_table(table, widths_dxa, header_fill=header_fill, font_size=font_size)
    doc.add_paragraph()
    return table


def add_note_box(doc, text):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = text
    format_table(table, [9360], header_fill=CALLOUT_FILL, font_size=9.4)
    shade_cell(cell, CALLOUT_FILL)
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in paragraph.runs:
            set_run_font(run, size=9.4, color=ACCENT)
    doc.add_paragraph()


def append_config_section(doc):
    doc.add_page_break()
    add_heading(doc, "4.3 Config 界面介绍", "SOP Heading 2")
    add_body_paragraph(
        doc,
        "概述：Config 界面用于配置 OHB 设备相关运行参数，包括 Idle Purge、气控阀压力、SH85 周期自检、湿度补偿、Purge Flow 和设备启用状态等。进入该界面前需确认当前账号具有 Config 权限。",
        after=8,
    )
    add_body_paragraph(doc, "Config 界面整体如下：", after=4, keep_next=True)
    add_image(doc, IMG_CONFIG_OVERVIEW, width_inches=6.2)
    add_caption(doc, "图 4-13 Config 界面整体布局与子标签导航（正文中的【1】、【2】对应本图编号）")

    add_heading(doc, "4.3.1 导航栏与子标签说明", "SOP Heading 3")
    add_body_paragraph(
        doc,
        "如图 4-13 所示，Config 界面顶部【1】为子标签导航栏，用于切换不同配置页面；各配置区域顶部【2】为当前配置模块的标题标签，用于提示当前正在操作的具体功能。",
        after=4,
        keep_next=True,
    )
    add_table(
        doc,
        ("子标签", "功能说明"),
        [
            ("Idle Purge", "配置所有设备的 Idle Purge 功能，例如准备时间、启用状态、充气持续时间和循环间隔。"),
            ("Pneumatic Valve", "配置气控阀主设备的目标压力，可按单个二维码设备设置，也可对所有可控主设备统一设置。"),
            ("SH85 Periodic", "配置 SH85 周期自检功能，例如自检启用状态、自检周期、自检状态显示和自检报告入口。"),
            ("SH85 Self-check", "用于手动触发单个设备执行一次 SH85 湿度传感器自检，可选择目标设备 ID 后点击 Check。"),
            ("Humidity Offset", "用于湿度采集值的偏移或补偿配置；配置前应确认补偿依据和现场校准要求。"),
            ("Purge Flow", "用于配置 VEFC 充氮流量，可按单个目标设备设置，也可对所有目标设备统一设置。"),
            ("Device Enable", "用于配置设备 Enable / Disable 状态。Disable 后设备状态仍可显示，但 Idle Purge、SH85 自检和充氮功能不再起作用。"),
        ],
        [2600, 6760],
    )

    add_heading(doc, "4.3.2 Idle Purge Configuration", "SOP Heading 3")
    add_body_paragraph(
        doc,
        "Idle Purge Configuration 用于配置所有设备的 Idle Purge 运行参数。该配置会影响设备在空闲状态下的自动充气行为，修改前应确认现场工艺要求。",
        after=4,
        keep_next=True,
    )
    add_image(doc, IMG_IDLE_PURGE, width_inches=6.25)
    add_caption(doc, "图 4-14 Idle Purge Configuration 区域与编号标注")
    add_table(
        doc,
        ("编号", "配置项", "功能说明"),
        [
            ("1", "Idle Purge Configuration", "Idle Purge 配置功能标题，用于说明当前区域为 Idle Purge 参数配置。"),
            ("2", "Preparation Time", "功能准备阶段时间，仅用于显示固定准备阶段时长，当前为只读项。Idle Purge 开始前会先经过该准备阶段。"),
            ("3", "Idle Purge Enable", "用于启用或禁用所有设备的 Idle Purge 功能。选择 Enable / Disable 后，点击 Set 写入配置。"),
            ("4", "Purge Duration", "用于设置 Idle Purge 充气状态持续时间，单位为秒。该时间决定每轮充气动作持续多久。"),
            ("5", "Purge Interval", "用于设置 Idle Purge 空闲状态持续时间，单位为秒。该时间决定两轮 Idle Purge 之间的等待间隔。"),
        ],
        [780, 2500, 6080],
    )

    add_heading(doc, "4.3.3 Pneumatic Valve Pressure Configuration", "SOP Heading 3")
    add_body_paragraph(
        doc,
        "Pneumatic Valve Pressure Configuration 用于配置气控阀压力。该功能支持按目标二维码设备单独设置，也支持对所有可控气控阀主设备进行统一设置。",
        after=4,
        keep_next=True,
    )
    add_image(doc, IMG_PNEUMATIC, width_inches=6.25)
    add_caption(doc, "图 4-15 Pneumatic Valve Pressure Configuration 区域与编号标注")
    add_table(
        doc,
        ("编号", "配置项", "功能说明"),
        [
            ("1", "Pneumatic Valve Pressure Configuration", "气控阀压力配置功能标题，用于说明当前区域为气控阀压力设置。"),
            ("2", "Target Device", "目标设备选择框。Combo Box 中包含所有可以控制气控阀的主设备 QRCode，选择后可对该设备进行单独设置。"),
            ("3", "Pneumatic Valve Pressure", "设置气控阀压力，单位为 bar。点击 Set 写入当前目标设备；点击 Set All 将该压力统一写入所有可控气控阀主设备。界面提示寄存器值 = pressure x 10000。"),
        ],
        [780, 2650, 5930],
    )

    add_heading(doc, "4.3.4 SH85 Periodic Self-check Configuration", "SOP Heading 3")
    add_body_paragraph(
        doc,
        "SH85 Periodic Self-check Configuration 用于配置 SH85 周期自检。该功能按设定周期自动执行自检，并在界面中显示当前自检状态、倒计时或本轮自检耗时。",
        after=4,
        keep_next=True,
    )
    add_image(doc, IMG_PERIODIC, width_inches=6.25)
    add_caption(doc, "图 4-16 SH85 Periodic Self-check Configuration 区域与编号标注")
    add_table(
        doc,
        ("编号", "配置项", "功能说明"),
        [
            ("1", "SH85 Periodic Self-check Configuration", "SH85 周期自检配置功能标题。"),
            ("2", "Enable Periodic Self-check", "启用或禁用 SH85 周期自检功能，选择 true / false 后写入配置。"),
            ("3", "Self-check Period", "设置两轮周期自检间隔；选择数值和单位后点击 Set。"),
            ("4", "Self-check Status", "显示当前周期自检状态。空闲状态下实时显示下一轮自检倒计时；自检状态下实时显示当前自检已经花费的时间。"),
            ("5", "Self-check Report", "点击 Open Report 按钮后弹出 SH85 Periodic Self-check Report 自检报告界面，可查看实时报告和历史报告。"),
        ],
        [780, 2750, 5830],
    )
    add_body_paragraph(doc, "固件侧自检状态 / 结果码说明如下：", after=4, keep_next=True)
    add_table(
        doc,
        ("状态 / 结果码", "含义"),
        [
            ("0", "无报警 / 空闲"),
            ("1", "自检进行中"),
            ("2", "自检成功"),
            ("3", "湿度超标失败"),
            ("4", "传感器通讯故障"),
            ("5", "值参数错误"),
        ],
        [1900, 7460],
        header_fill=STATUS_FILL,
    )
    add_body_paragraph(doc, "固件自检流程如下：", after=4, keep_next=True)
    add_image(doc, IMG_FLOW, width_inches=2.65)
    add_caption(doc, "图 4-17 SH85 周期自检固件流程图")
    add_note_box(
        doc,
        "流程说明：启动 RH Sensor 自检后，固件会先保存所有阀状态，再自动关闭所有气路、关闭 VEFC 并打开自检气路阀；持续充气 60s 后采集湿度值并与标准值比较，随后返回自检结果 / 状态并恢复阀位。若达到标准值则结束；若未达到标准值，则需要返厂标定。"
    )

    add_heading(doc, "4.3.5 SH85 Periodic Self-check Report", "SOP Heading 3")
    add_body_paragraph(
        doc,
        "在图 4-16 中点击【5】Open Report 后，会打开 SH85 Periodic Self-check Report 自检报告窗口。报告窗口包含 Live Log 和 History Log 两个子标签。",
        after=4,
        keep_next=True,
    )
    add_body_paragraph(doc, "Live Log 实时报告界面如下：", after=4, keep_next=True)
    add_image(doc, IMG_REPORT_LIVE, width_inches=6.15)
    add_caption(doc, "图 4-18 SH85 周期自检实时报告界面（Live Log）")
    add_table(
        doc,
        ("字段", "功能说明"),
        [
            ("QRCode", "参与自检的设备二维码编号，用于定位具体设备。"),
            ("Execution Status", "显示设备当前自检执行状态或阶段，例如等待下一阶段、执行中或未执行等。"),
            ("Countdown(s)", "显示当前阶段剩余倒计时，单位为秒；无有效倒计时时显示为 -。"),
            ("Success", "显示当前轮自检是否成功；自检尚未完成时通常显示为 -。"),
            ("Participated", "显示该设备是否参与当前轮周期自检，Yes 表示参与。"),
        ],
        [2300, 7060],
    )
    add_body_paragraph(doc, "History Log 历史报告界面如下：", after=4, keep_next=True)
    add_image(doc, IMG_REPORT_HISTORY, width_inches=6.15)
    add_caption(doc, "图 4-19 SH85 周期自检历史报告界面（History Log）")
    add_table(
        doc,
        ("字段", "功能说明"),
        [
            ("Last Check Start Time", "上一轮自检开始时间，用于追溯历史自检发生时间。"),
            ("Success Count", "上一轮自检成功数量。"),
            ("Failure Count", "上一轮自检失败数量。"),
            ("Participated", "显示设备是否参与上一轮自检，Yes 表示参与。"),
            ("Description", "显示上一轮自检结果信息，例如湿度超标、设备未连接或 SH85 传感器通讯错误等。失败信息会在界面中突出显示。"),
        ],
        [2450, 1450, 1450, 1450, 4560],
        font_size=8.7,
    )

    add_heading(doc, "4.3.6 SH85 Self-check Configuration", "SOP Heading 3")
    add_body_paragraph(
        doc,
        "SH85 Self-check Configuration 用于手动触发 SH85 湿度传感器自检。该功能只会对当前选定的一个设备执行一次自检，不会对所有设备批量执行。",
        after=4,
        keep_next=True,
    )
    add_image(doc, IMG_MANUAL_SELF_CHECK, width_inches=6.25)
    add_caption(doc, "图 4-20 SH85 Self-check Configuration 区域与编号标注")
    add_table(
        doc,
        ("编号", "配置项", "功能说明"),
        [
            ("1", "SH85 Self-check Configuration", "手动 SH85 自检功能标题，用于说明当前区域为手动自检操作区。"),
            ("2", "Target Device ID", "选择需要执行手动自检的目标设备 ID。该功能仅对当前选中的单个设备生效。"),
            ("3", "SH85 Self-check", "点击 Check 按钮后触发一次 SH85 湿度传感器自检，界面提示单次自检约 70 秒。"),
        ],
        [780, 2600, 5980],
    )

    add_heading(doc, "4.3.7 Purge Flow Configuration", "SOP Heading 3")
    add_body_paragraph(
        doc,
        "Purge Flow Configuration 用于配置 VEFC 充氮流量。该参数只在 FOUP IN 状态下有效，修改前应确认现场流量设定规范。",
        after=4,
        keep_next=True,
    )
    add_image(doc, IMG_PURGE_FLOW, width_inches=6.25)
    add_caption(doc, "图 4-21 Purge Flow Configuration 区域与编号标注")
    add_table(
        doc,
        ("编号", "配置项", "功能说明"),
        [
            ("1", "Purge Flow Configuration", "充氮流量配置功能标题，用于说明当前区域为 Purge Flow 设置。"),
            ("2", "Target Device", "目标设备选择框，用于选择需要设置 Purge Flow 的设备 QRCode。"),
            ("3", "Purge Flow", "设置 VEFC 充氮流量，单位为 L/Min。点击 Set 写入当前目标设备；点击 Set All 可统一写入所有目标设备。界面提示寄存器值 = flow x 100。"),
        ],
        [780, 2600, 5980],
    )

    doc.add_page_break()
    add_heading(doc, "4.3.8 Device Enable Configuration", "SOP Heading 3")
    add_body_paragraph(
        doc,
        "Device Enable Configuration 用于配置设备是否启用。该配置会影响设备是否参与自动功能流程，修改前应确认设备当前状态和现场操作要求。",
        after=4,
        keep_next=True,
    )
    add_image(doc, IMG_DEVICE_ENABLE, width_inches=6.25)
    add_caption(doc, "图 4-22 Device Enable Configuration 区域与编号标注")
    add_table(
        doc,
        ("编号", "配置项", "功能说明"),
        [
            ("1", "Device Enable Configuration", "设备使能配置功能标题，用于说明当前区域为设备 Enable / Disable 设置。"),
            ("2", "Target Device", "目标设备选择框，用于选择需要设置使能状态的设备 QRCode。"),
            ("3", "Device Status", "设置设备状态为 Enable 或 Disable，选择后点击 Set 写入。"),
        ],
        [780, 2600, 5980],
    )
    add_note_box(
        doc,
        "Disable 状态说明：当设备处于 Disable 状态时，设备的温度、湿度等状态数据仍能够正常显示；但 Idle Purge、SH85 自检和充氮功能不再起作用。"
    )


def main():
    for image in (
        IMG_CONFIG_OVERVIEW,
        IMG_IDLE_PURGE,
        IMG_PNEUMATIC,
        IMG_PERIODIC,
        IMG_FLOW,
        IMG_REPORT_LIVE,
        IMG_REPORT_HISTORY,
        IMG_MANUAL_SELF_CHECK,
        IMG_PURGE_FLOW,
        IMG_DEVICE_ENABLE,
    ):
        if not image.exists():
            raise FileNotFoundError(image)

    doc = Document(SOURCE)
    append_config_section(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()

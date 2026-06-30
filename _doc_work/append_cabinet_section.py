from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SOURCE = Path(r"C:\Users\Lenovo\Desktop\OHB80PortMonitor软件SOP_格式优化版.docx")
IMAGE = Path(r"C:\Users\Lenovo\AppData\Local\Temp\codex-clipboard-db2394fc-d7a7-4237-97ca-301d1aa4c057.png")
WORK_OUTPUT = Path(r"D:\Project\CYTC_Project\OHB80PortMonitor\_doc_work\OHB80PortMonitor_SOP_add_cabinet.docx")

BODY_FONT = "Microsoft YaHei"
ACCENT = RGBColor(31, 77, 120)
HEADING = RGBColor(46, 116, 181)
MUTED = RGBColor(89, 89, 89)
INK = RGBColor(33, 33, 33)
TABLE_HEADER_FILL = "E8EEF5"
TABLE_BORDER = "B7C6D8"
CALLOUT_FILL = "F4F6F9"


def set_run_font(run, name=BODY_FONT, size=None, bold=None, color=None):
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:eastAsia"), name)
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def clear_paragraph(paragraph):
    p = paragraph._element
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.25):
    fmt = paragraph.paragraph_format
    fmt.left_indent = Pt(0)
    fmt.right_indent = Pt(0)
    fmt.first_line_indent = Pt(0)
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_field(paragraph, instr_text):
    run_begin = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run_begin._r.append(fld_char_begin)

    run_instr = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instr_text
    run_instr._r.append(instr)

    run_sep = paragraph.add_run()
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    run_sep._r.append(fld_char_sep)

    run_text = paragraph.add_run("1")
    set_run_font(run_text, size=9, color=MUTED)

    run_end = paragraph.add_run()
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run_end._r.append(fld_char_end)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=140, bottom=90, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "6")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), TABLE_BORDER)


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_table_geometry(table, widths_dxa):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            if index < len(widths_dxa):
                set_cell_width(cell, widths_dxa[index])


def style_paragraph(paragraph, style_name=None, before=0, after=6, line=1.25, color=INK, size=10.5, bold=False):
    if style_name is not None:
        paragraph.style = style_name
    set_paragraph_spacing(paragraph, before=before, after=after, line=line)
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=bold, color=color)


def update_headers(doc):
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        clear_paragraph(paragraph)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(paragraph, before=0, after=2, line=1.0)
        run = paragraph.add_run("OHB80PortMonitor 软件 SOP | 界面操作说明")
        set_run_font(run, size=9, color=MUTED)


def ensure_footer(doc):
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        if paragraph.text.strip():
            continue
        clear_paragraph(paragraph)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(paragraph, before=0, after=0, line=1.0)
        run = paragraph.add_run("第 ")
        set_run_font(run, size=9, color=MUTED)
        add_field(paragraph, "PAGE")
        run = paragraph.add_run(" 页")
        set_run_font(run, size=9, color=MUTED)


def add_body_paragraph(doc, text, before=0, after=6, keep_next=False):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5, color=INK)
    set_paragraph_spacing(paragraph, before=before, after=after, line=1.25)
    paragraph.paragraph_format.keep_with_next = keep_next
    return paragraph


def add_caption(doc, text="如下："):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    set_run_font(run, size=9, color=MUTED)
    set_paragraph_spacing(paragraph, before=2, after=4, line=1.15)
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_heading(doc, text, style_name, before=None):
    paragraph = doc.add_paragraph(text)
    paragraph.style = style_name
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if before is not None:
        paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.keep_with_next = True
    for run in paragraph.runs:
        if style_name.endswith("2"):
            set_run_font(run, size=14, bold=True, color=HEADING)
        else:
            set_run_font(run, size=12, bold=True, color=ACCENT)
    return paragraph


def format_table(table, widths_dxa, header_fill=TABLE_HEADER_FILL):
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.autofit = False
    set_table_borders(table)
    set_table_geometry(table, widths_dxa)
    if table.rows:
        set_repeat_header(table.rows[0])
    for row_index, row in enumerate(table.rows):
        for col_index, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            if row_index == 0:
                shade_cell(cell, header_fill)
            for paragraph in cell.paragraphs:
                set_paragraph_spacing(paragraph, before=0, after=0, line=1.18)
                if row_index == 0 or col_index == 0:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    set_run_font(run, size=9.3 if row_index else 9.5, bold=(row_index == 0), color=INK)


def add_area_table(doc):
    rows = [
        ("1", "控制按钮状态（Switch Control）", "显示电控柜各路控制对象的当前开关状态，包括 Inlet Fan、Exhaust Fan、Red Light、Green Light 和 Control Power。状态列显示 On / Off，用于判断当前输出是否启用。"),
        ("2", "串口连接状态（Cabinet Connection）", "显示上位机与电控柜控制板 / 串口通讯连接状态。Connected 表示通讯正常，可进行状态读取和控制操作。"),
        ("3", "按钮控制", "用于手动控制电控柜相关输出。包含 EMO、Control Power、Red Light、Green Light、Inlet Fan、Exhaust Fan、Fan Master 等按钮；执行前应确认当前用户权限、连接状态和现场设备状态。"),
        ("4", "警报状态（Alarm Status）", "显示 EMO1、EMO2、Smoke Alarm 等安全 / 联锁状态。Interlock Released 表示联锁释放；Normal 表示状态正常。"),
        ("5", "电源 / 电流（Voltage / Current）", "显示电控柜电源与电流采集值，例如 Phase A Voltage、Phase A Current，用于观察供电是否处于正常范围。"),
        ("6", "温度 / 湿度（Temp / Humidity）", "显示电控柜温湿度采集值，并提供温度、湿度设定入口，可通过 Set Temp (C) 和 Set Humidity (%) 写入目标设定值。"),
    ]
    table = doc.add_table(rows=1, cols=3)
    headers = ("编号", "区域", "功能说明")
    for index, text in enumerate(headers):
        table.cell(0, index).text = text
    for row in rows:
        cells = table.add_row().cells
        for index, text in enumerate(row):
            cells[index].text = text
    format_table(table, [780, 3300, 5280])
    doc.add_paragraph()


def add_control_table(doc):
    rows = [
        ("Control Power", "控制 OHB 设备 / 电控柜相关电源输出。操作前需确认现场设备允许上电或断电。"),
        ("Fan Master", "风扇总使能控制；通常作为 Inlet Fan、Exhaust Fan 控制前的总开关条件。"),
        ("Inlet Fan / Exhaust Fan", "分别控制进风扇和排风扇，用于电控柜通风散热。"),
        ("Red Light / Green Light", "控制红灯、绿灯指示输出，用于现场状态提示。"),
        ("EMO", "用于执行报警取消或解除相关状态。操作前需确认安全条件和现场 SOP。"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "控制项"
    table.cell(0, 1).text = "说明"
    for row in rows:
        cells = table.add_row().cells
        cells[0].text = row[0]
        cells[1].text = row[1]
    format_table(table, [2450, 6910])
    doc.add_paragraph()


def add_note_box(doc, text):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = text
    shade_cell(cell, CALLOUT_FILL)
    set_cell_margins(cell, top=120, start=160, bottom=120, end=160)
    format_table(table, [9360], header_fill=CALLOUT_FILL)
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in paragraph.runs:
            set_run_font(run, size=9.5, color=ACCENT, bold=False)
    doc.add_paragraph()


def append_cabinet_section(doc):
    doc.add_page_break()
    add_heading(doc, "4.2 Cabinet 界面介绍", "SOP Heading 2")
    add_body_paragraph(
        doc,
        "概述：Cabinet 界面用于监控电控柜状态，并提供电控柜相关控制功能，包括控制 OHB 设备电源、控制风扇、控制指示灯和取消报警等。",
        after=8,
    )
    add_body_paragraph(doc, "Cabinet 界面整体如下：", after=4, keep_next=True)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(IMAGE), width=Inches(6.15))
    set_paragraph_spacing(paragraph, before=2, after=10, line=1.0)

    add_heading(doc, "4.2.1 界面区域说明", "SOP Heading 3")
    add_body_paragraph(doc, "根据图中编号，Cabinet 界面主要区域说明如下：", after=4, keep_next=True)
    add_area_table(doc)

    add_heading(doc, "4.2.2 控制功能说明", "SOP Heading 3")
    add_body_paragraph(doc, "按钮控制区域用于执行电控柜相关输出控制。常用控制项说明如下：", after=4, keep_next=True)
    add_control_table(doc)

    add_heading(doc, "4.2.3 操作注意事项", "SOP Heading 3")
    add_note_box(
        doc,
        "注意：执行电源、风扇、指示灯或报警取消操作前，应先确认 Cabinet Connection 为 Connected，并确认当前登录账号具备对应权限。若警报状态未恢复正常，应优先排查现场设备与安全互锁状态，避免直接重复操作。",
    )


def main():
    doc = Document(SOURCE)
    update_headers(doc)
    ensure_footer(doc)
    append_cabinet_section(doc)
    doc.save(WORK_OUTPUT)
    print(WORK_OUTPUT)


if __name__ == "__main__":
    main()

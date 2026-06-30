# -*- coding: utf-8 -*-
from pathlib import Path
import re

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SOURCE = Path(r"D:\Project\CYTC_Project\OHB80PortMonitor\_doc_work\OHB80PortMonitor_SOP_add_config.docx")
OUTPUT = Path(r"D:\Project\CYTC_Project\OHB80PortMonitor\_doc_work\OHB80PortMonitor_SOP_add_config_home_ohb_ui_runlog.docx")

IMG_OPERATION_LIVE = Path(r"C:\Users\Lenovo\AppData\Local\Temp\codex-clipboard-6404c62d-088b-467c-9b0f-a17a1a4a9476.png")
IMG_OPERATION_HISTORY = Path(r"C:\Users\Lenovo\AppData\Local\Temp\codex-clipboard-35db89e1-80bc-476f-9f8c-9ad9da66a90b.png")
IMG_OPERATION_TIME_DIALOG = Path(r"C:\Users\Lenovo\AppData\Local\Temp\codex-clipboard-09c4b99a-6858-4e20-b486-ce2c0f8372d9.png")
IMG_OPERATION_SEARCH_RESULT = Path(r"C:\Users\Lenovo\AppData\Local\Temp\codex-clipboard-ecd8e9d6-6424-4b8b-9647-23fe12fd756d.png")
IMG_SOFTWARE_UI = Path(r"C:\Users\Lenovo\AppData\Local\Temp\codex-clipboard-4ce93899-193e-41eb-bf7b-099a4f2de95e.png")
IMG_FIRMWARE_FLOW = Path(r"C:\Users\Lenovo\AppData\Local\Temp\codex-clipboard-c0b17364-cdca-4a8b-bac9-b720b445134f.png")
IMG_FIRMWARE_LOGO = Path(r"C:\Users\Lenovo\AppData\Local\Temp\codex-clipboard-470be2ae-5386-41f0-9133-b5732738a2fb.png")
IMG_FIRMWARE_PAGE1 = Path(r"C:\Users\Lenovo\AppData\Local\Temp\codex-clipboard-384f1f9d-5ecd-4216-8169-ca27cf036229.png")
IMG_FIRMWARE_PAGE2 = Path(r"C:\Users\Lenovo\AppData\Local\Temp\codex-clipboard-8264f5fd-490d-4b0e-bcac-037c63969350.png")

BODY_FONT = "Microsoft YaHei"
ACCENT = RGBColor(31, 77, 120)
HEADING = RGBColor(46, 116, 181)
MUTED = RGBColor(89, 89, 89)
INK = RGBColor(33, 33, 33)
TABLE_HEADER_FILL = "E8EEF5"
TABLE_BORDER = "B7C6D8"
CALLOUT_FILL = "F4F6F9"


def set_run_font(run, name=BODY_FONT, size=None, bold=None, color=None):
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:eastAsia"), name)
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.25):
    fmt = paragraph.paragraph_format
    fmt.left_indent = Pt(0)
    fmt.right_indent = Pt(0)
    fmt.first_line_indent = Pt(0)
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=140, bottom=90, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "6")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), TABLE_BORDER)


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def set_table_geometry(table, widths_dxa):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            if index < len(widths_dxa):
                set_cell_width(cell, widths_dxa[index])


def add_heading(doc, text, style_name, before=None):
    paragraph = doc.add_paragraph(text)
    paragraph.style = style_name
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.keep_with_next = True
    if before is not None:
        paragraph.paragraph_format.space_before = Pt(before)
    for run in paragraph.runs:
        if style_name.endswith("2"):
            set_run_font(run, size=14, bold=True, color=HEADING)
        else:
            set_run_font(run, size=12, bold=True, color=ACCENT)
    return paragraph


def add_body_paragraph(doc, text, before=0, after=6, keep_next=False):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5, color=INK)
    set_paragraph_spacing(paragraph, before=before, after=after, line=1.25)
    paragraph.paragraph_format.keep_with_next = keep_next
    return paragraph


def add_label_paragraph(doc, text):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5, bold=True, color=ACCENT)
    set_paragraph_spacing(paragraph, before=2, after=4, line=1.15)
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_image(doc, image_path, width_inches=6.15):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    set_paragraph_spacing(paragraph, before=2, after=2, line=1.0)
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    return paragraph


def add_caption(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(paragraph, before=2, after=8, line=1.0)
    run = paragraph.add_run(text)
    set_run_font(run, size=9, color=MUTED)
    return paragraph


def format_table(table, widths_dxa, header_fill=TABLE_HEADER_FILL, font_size=9.2):
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.autofit = False
    set_table_borders(table)
    set_table_geometry(table, widths_dxa)
    if table.rows:
        set_repeat_header(table.rows[0])
    for row_index, row in enumerate(table.rows):
        set_cant_split(row)
        for col_index, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            if row_index == 0:
                shade_cell(cell, header_fill)
            for paragraph in cell.paragraphs:
                set_paragraph_spacing(paragraph, before=0, after=0, line=1.18)
                if row_index == 0 or col_index == 0:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    set_run_font(run, size=font_size if row_index else 9.5, bold=(row_index == 0), color=INK)


def add_table(doc, headers, rows, widths_dxa, header_fill=TABLE_HEADER_FILL, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    for index, text in enumerate(headers):
        table.cell(0, index).text = text
    for row in rows:
        cells = table.add_row().cells
        for index, text in enumerate(row):
            cells[index].text = text
    format_table(table, widths_dxa, header_fill=header_fill, font_size=font_size)
    doc.add_paragraph()
    return table


def add_note_box(doc, text):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = text
    format_table(table, [9360], header_fill=CALLOUT_FILL, font_size=9.3)
    shade_cell(cell, CALLOUT_FILL)
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in paragraph.runs:
            set_run_font(run, size=9.3, color=ACCENT)
    doc.add_paragraph()


def replace_in_paragraph(paragraph, pattern, repl):
    text = paragraph.text
    new_text = pattern.sub(repl, text)
    if new_text == text:
        return
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = new_text


def renumber_following_figures(doc):
    pattern = re.compile(r"图 4-(\d+)")

    def repl(match):
        number = int(match.group(1))
        if number >= 12:
            return f"图 4-{number + 9}"
        if number >= 7:
            return f"图 4-{number + 4}"
        return match.group(0)

    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, pattern, repl)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph, pattern, repl)


def append_operation_log_detail_section(doc):
    add_label_paragraph(doc, "运行日志界面详细说明")
    add_body_paragraph(
        doc,
        "运行日志界面由 Live Log 和 History 两个子标签组成。Live Log 用于查看当前软件运行过程中的实时日志，History 用于按日志类型、记录时间和关键字查询数据库中的历史日志记录。",
        after=8,
    )

    add_label_paragraph(doc, "Live Log 实时日志界面")
    add_body_paragraph(doc, "点击运行日志公告栏后，默认进入 Live Log 实时日志界面，如下图所示：", after=4, keep_next=True)
    add_image(doc, IMG_OPERATION_LIVE, width_inches=5.85)
    add_caption(doc, "图 4-7 运行日志实时记录界面与字段说明")
    add_table(
        doc,
        ("编号", "区域", "说明"),
        [
            ("1", "运行日志导航栏", "用于在 Live Log 和 History 两个子标签之间切换。Live Log 显示实时日志，History 显示历史查询界面。"),
            ("2", "日志表头", "实时日志表格的字段标题区，包含 Occur Time、Log Type、Description 三个字段。"),
        ],
        [780, 2300, 6280],
    )
    add_table(
        doc,
        ("字段", "含义", "说明"),
        [
            ("Occur Time", "日志发生时间", "显示该条日志产生的具体时间，用于追溯事件发生顺序。"),
            ("Log Type", "日志类型", "用于区分日志级别或类型，常见类型包括 Message 和 Error。Error 类型通常以红色文字提示。"),
            ("Description", "日志描述", "显示具体日志内容，例如设备离线、报警解除、命令执行结果或 SH85 自检结果等。"),
        ],
        [1700, 2300, 5360],
    )
    add_note_box(
        doc,
        "稳定性说明：Live Log 只作为实时可见窗口使用，界面侧最多保留最近 2000 条可见日志。新日志会先根据当前登录用户权限判断是否可见，然后增量追加到界面；超过限制时批量裁剪旧记录，避免界面模型无限增长。"
    )

    add_label_paragraph(doc, "History 历史日志界面")
    add_body_paragraph(doc, "点击 History 子标签后，进入运行日志历史查询界面，如下图所示：", after=4, keep_next=True)
    add_image(doc, IMG_OPERATION_HISTORY, width_inches=5.85)
    add_caption(doc, "图 4-8 运行日志历史记录查询界面与查询条件")
    add_table(
        doc,
        ("编号", "区域", "说明"),
        [
            ("1", "所有条件都启用", "用于快速启用查询条件。启用后，界面按当前配置的日志类型、记录时间和关键字等条件组合查询。"),
            ("2", "日志类型条件", "勾选后可通过 Log Type 下拉框选择日志类型，例如 Message 或 Error，用于缩小查询范围。"),
            ("3", "时间范围条件", "勾选 Record Time 后启用时间范围过滤；点击 Set 按钮会弹出时间设置弹框。"),
            ("4", "关键字查询条件", "在输入框中输入需要查询的关键字，点击 Search 按钮即可搜索包含该关键字的记录。"),
            ("5", "分页按钮", "用于切换历史日志结果页，包括上一页、页码、下一页、当前页/总页数和跳转页码输入框。"),
        ],
        [780, 2200, 6380],
    )
    add_table(
        doc,
        ("字段", "含义", "说明"),
        [
            ("Occur Time", "日志发生时间", "显示历史日志记录产生的时间。"),
            ("Log Type", "日志类型", "显示该条记录的类型，便于按消息、错误等类别筛选。"),
            ("Description", "日志描述", "显示完整日志内容，关键字查询会在该字段及相关文本中进行匹配。"),
        ],
        [1700, 2300, 5360],
    )

    add_label_paragraph(doc, "时间范围设置")
    add_body_paragraph(
        doc,
        "点击图 4-8 中【3】区域的 Set 按钮后，会弹出 Set Record Time 时间设置弹框，如下图所示：",
        after=4,
        keep_next=True,
    )
    add_image(doc, IMG_OPERATION_TIME_DIALOG, width_inches=4.65)
    add_caption(doc, "图 4-9 运行日志时间范围设置弹框")
    add_table(
        doc,
        ("字段 / 控件", "说明"),
        [
            ("Enable Start Time", "勾选后启用开始时间过滤，查询结果只保留晚于或等于开始时间的日志。"),
            ("Enable End Time", "勾选后启用结束时间过滤，查询结果只保留早于或等于结束时间的日志。"),
            ("日期 / 时间输入框", "分别设置开始时间和结束时间的日期与具体时分秒。"),
            ("OK", "确认当前时间范围设置，并返回 History 查询界面。"),
            ("Cancel", "取消本次设置，不修改当前时间范围条件。"),
        ],
        [2600, 6760],
    )

    add_label_paragraph(doc, "关键字搜索与记录定位")
    add_body_paragraph(
        doc,
        "在 History 界面中，关键字搜索用于从历史日志中定位包含指定内容的记录。输入关键字后点击 Search，界面会显示查询结果，并可通过 Pre、Next、Record No. 和 Jump 在匹配记录之间快速切换。",
        after=4,
        keep_next=True,
    )
    add_table(
        doc,
        ("控件", "功能说明"),
        [
            ("Search", "按当前启用的查询条件和关键字执行搜索。搜索完成后，第一条匹配记录会被自动选中。"),
            ("Pre", "跳转到上一条匹配记录；如果上一条记录位于其他页面，界面会自动切换到对应页面并选中该记录。位于第一条时继续点击会循环跳转到最后一条。"),
            ("Next", "跳转到下一条匹配记录；如果下一条记录位于其他页面，界面会自动切换到对应页面并选中该记录。位于最后一条时继续点击会循环跳转到第一条。"),
            ("0/0", "显示当前选中匹配记录序号 / 匹配记录总数。无关键字或无匹配结果时显示 0/0。"),
            ("Record No.", "输入希望直接查看的匹配记录序号。这里的序号表示关键字匹配结果中的第几条记录。"),
            ("Jump", "根据 Record No. 输入的序号直接跳转到对应记录；如目标记录不在当前页，界面会自动切换到目标记录所在页面。"),
        ],
        [1900, 7460],
        font_size=8.9,
    )
    add_body_paragraph(doc, "下面以关键字 12001 为例说明查询结果的显示方式：", after=4, keep_next=True)
    add_image(doc, IMG_OPERATION_SEARCH_RESULT, width_inches=5.85)
    add_caption(doc, "图 4-10 关键字 12001 查询结果与匹配记录定位示例")
    add_table(
        doc,
        ("编号", "区域", "说明"),
        [
            ("1", "条件查询被选中的第一条记录", "输入 12001 并点击 Search 后，界面会自动定位到第一条匹配记录，该记录使用黄色边框或高亮表示当前选中状态。"),
            ("2", "查询到的其他匹配记录", "同一查询结果中的其他匹配记录会以绿色背景提示，便于用户识别当前页内还存在其他包含关键字的日志。"),
        ],
        [780, 2650, 5930],
    )
    add_note_box(
        doc,
        "操作提示：History 的分页按钮用于切换当前结果页；Pre、Next 和 Jump 用于在关键字匹配记录之间定位。两类操作可以配合使用，但当关键字匹配记录跨页时，Pre、Next 和 Jump 会自动完成页面切换。"
    )
    doc.add_page_break()


def append_home_ohb_ui_section(doc):
    add_heading(doc, "4.1.9 软件 OHB 设备 UI 与固件 OHB 屏幕 UI", "SOP Heading 3")
    add_body_paragraph(
        doc,
        "本节用于说明软件 Home 视图中的 OHB 设备卡片、上方详细信息区域，以及固件端 OHB 屏幕三页 UI 之间的对应关系。客户可通过本节快速判断软件显示内容与现场设备屏幕显示内容是否一致。",
        after=8,
    )

    add_label_paragraph(doc, "软件 OHB 设备 UI")
    add_body_paragraph(doc, "软件端 OHB 设备 UI 如下图所示：", after=4, keep_next=True)
    add_image(doc, IMG_SOFTWARE_UI, width_inches=6.25)
    add_caption(doc, "图 4-16 软件 OHB 设备 UI 与编号标注")
    add_table(
        doc,
        ("编号", "区域", "说明"),
        [
            ("1", "设备在视图上的信息", "设备卡片仅显示 ID(QRCode)、P(进气压力)、F(进气流量)、RH(相对湿度)，用于在 Home 主视图中快速查看单台 OHB 设备的关键状态。"),
            ("2", "设备具体信息", "选中设备后，上方 Detailed Information 区域会显示该设备的完整实时数据和 Idle Purge 状态，字段含义见下表。"),
            ("3", "设备背景颜色", "设备背景颜色用于提示设备当前状态。该规则已在 4.1.6 设备背景颜色中说明，包括充氮时长、FoupOut、Alarm、Disable 等状态，客户可回顾 4.1.6 进行对照。"),
        ],
        [780, 2350, 6230],
    )
    add_table(
        doc,
        ("字段", "含义", "与设备卡片/固件屏幕的对应关系"),
        [
            ("QRCode", "设备二维码或设备 ID。", "对应软件设备卡片中的 ID，也对应固件屏幕第二页的 QRCode。"),
            ("InletPressure", "进气压力。", "对应软件设备卡片中的 P，也对应固件屏幕第一页 PI_Mpa。"),
            ("OutletPressures", "出气压力。", "对应固件屏幕第二页 PO_Kpa。"),
            ("InletFlow", "进气流量。", "对应软件设备卡片中的 F，也对应固件屏幕第一页 F L/Min。"),
            ("Relative Humidity", "相对湿度。", "对应软件设备卡片中的 RH，也对应固件屏幕第一页 RH %。"),
            ("Temperature", "设备温度。", "对应固件屏幕第二页 Temp。"),
            ("Start Time", "Idle Purge 或相关运行动作的开始时间。", "用于查看当前运行阶段起点。"),
            ("Duration", "当前运行阶段持续时间。", "用于判断设备在当前阶段已持续多久。"),
            ("Purge Time", "充氮累计或阶段时间。", "用于结合充氮进度条和设备背景颜色判断充氮状态。"),
            ("Idle Enable", "Idle Purge 功能是否启用。", "对应固件屏幕第二页 IDLE 相关状态。"),
            ("Idle State", "Idle Purge 当前状态。", "对应固件屏幕第二页 IDLE 相关状态。"),
            ("IdleTime", "Idle Purge 时间信息。", "对应固件屏幕第二页 IDLE 显示值。"),
        ],
        [1900, 3100, 4360],
        font_size=8.6,
    )

    add_label_paragraph(doc, "固件 OHB 屏幕 UI")
    add_body_paragraph(
        doc,
        "固件 OHB 屏幕 UI 分为三页：厂商 Logo 页面、设备第一页、设备第二页。设备第一页偏向显示 Foup IN / Foup OUT、充氮进度、进气流量、进气压力和相对湿度；设备第二页偏向显示 QRCode、出气压力、报警状态、温度和 Idle Purge 状态。",
        after=8,
    )
    add_body_paragraph(doc, "固件 UI 屏运行流程如下：", after=4, keep_next=True)
    add_image(doc, IMG_FIRMWARE_FLOW, width_inches=5.75)
    add_caption(doc, "图 4-17 固件 UI 屏运行流程图")
    add_note_box(
        doc,
        "流程说明：固件 UI 屏运行逻辑主要包含页面切换、FOUP 状态显示和报警状态显示三部分。页面切换支持默认轮播和上位机配置；FOUP IN / FOUP OUT 状态决定屏幕参数展示方式；报警状态用于在异常触发时通过屏幕左下角报警指示提醒用户。"
    )

    add_body_paragraph(doc, "1. 厂商 Logo 页面如下：", after=4, keep_next=True)
    add_image(doc, IMG_FIRMWARE_LOGO, width_inches=4.95)
    add_caption(doc, "图 4-18 固件 OHB 屏幕 UI 厂商 Logo 页面")
    add_note_box(doc, "说明：厂商 Logo 页面用于展示设备供应商信息、系统名称、反馈二维码、联系电话和网址等基础信息，主要用于设备启动或待机展示。")

    add_body_paragraph(doc, "2. 设备第一页如下：", after=4, keep_next=True)
    add_image(doc, IMG_FIRMWARE_PAGE1, width_inches=5.6)
    add_caption(doc, "图 4-19 固件 OHB 屏幕 UI 设备第一页")
    add_table(
        doc,
        ("编号", "字段 / 区域", "功能说明"),
        [
            ("1", "Foup IN", "表示设备当前处于 Foup IN 状态；其颜色状态与软件设备 UI 的背景颜色状态对应。"),
            ("2", "Foup OUT", "表示设备当前处于 Foup OUT 状态；其颜色状态与软件设备 UI 的背景颜色状态对应。"),
            ("3", "充氮进度条", "进度条会随充氮时间变化而变色，对应图 4-16 中【3】设备背景颜色。正常达到 30 分钟后对应软件浅绿色状态；如果 30 分钟未达标，则对应软件 Alarm 红色报警状态。"),
            ("4", "F L/Min", "进气流量，对应软件设备卡片中的 F，也对应 Detailed Information 中的 InletFlow。"),
            ("5", "PI_Mpa", "进气压力，对应软件设备卡片中的 P，也对应 Detailed Information 中的 InletPressure。"),
            ("6", "RH %", "相对湿度，对应软件设备卡片中的 RH，也对应 Detailed Information 中的 Relative Humidity。"),
        ],
        [780, 2250, 6330],
    )

    add_body_paragraph(doc, "3. 设备第二页如下：", after=4, keep_next=True)
    add_image(doc, IMG_FIRMWARE_PAGE2, width_inches=5.6)
    add_caption(doc, "图 4-20 固件 OHB 屏幕 UI 设备第二页")
    add_table(
        doc,
        ("编号", "字段 / 区域", "功能说明"),
        [
            ("1", "QRCode", "设备二维码或设备 ID，对应图 4-16 软件设备 UI 中【1】显示的 ID(QRCode)。"),
            ("2", "充氮进度条", "与设备第一页【3】作用一致，用于显示充氮进度并对应软件设备 UI 的背景颜色状态。"),
            ("3", "PO_Kpa", "出气压力，对应图 4-16 软件设备 UI 中【2】Detailed Information 的 OutletPressures 字段。"),
            ("4", "报警图标", "表示设备发生报警，对应软件设备 UI 背景颜色中的 Alarm 红色状态；具体背景颜色含义可回顾 4.1.6。"),
            ("5", "Temp", "设备温度，对应图 4-16 软件设备 UI 中【2】Detailed Information 的 Temperature 字段；相对湿度 RH 已在设备第一页显示。"),
            ("6", "IDLE", "Idle Purge 状态或时间信息，对应图 4-16 软件设备 UI 中【2】Detailed Information 的 Idle Enable、Idle State、IdleTime 等字段。"),
        ],
        [780, 2250, 6330],
    )

    doc.add_page_break()
    add_label_paragraph(doc, "固件 UI 屏报警显示")
    add_body_paragraph(
        doc,
        "设备运行过程中如出现异常，UI 屏在接收到报警信号后会进入报警显示状态。报警用于提醒现场人员当前设备存在异常，需要结合上位机告警信息、设备状态和现场 SOP 进行确认。",
        after=4,
        keep_next=True,
    )
    add_table(
        doc,
        ("报警触发项", "说明"),
        [
            ("VEFC / VFFC 异常", "充氮流量控制相关模块出现异常时，UI 屏触发报警提示。"),
            ("温湿度传感器未接入", "UI 屏或固件检测到温湿度传感器未连接、未接入或无法正常读取时，触发报警提示。"),
            ("FOUP IN 超时未达标", "设备处于 FOUP IN 状态并运行 30 分钟后，若湿度仍未达到设定要求，UI 屏触发报警提示。"),
            ("温湿度传感器异常", "温湿度传感器通讯异常、采集异常或状态异常时，UI 屏触发报警提示。"),
        ],
        [2700, 6660],
    )
    add_note_box(
        doc,
        "报警显示逻辑：当上述任一报警条件成立时，UI 屏进入报警状态，并通过屏幕左下角报警指示闪烁提示用户当前设备异常。异常条件解除后，报警状态自动取消，报警指示恢复为正常显示状态。"
    )

    add_label_paragraph(doc, "固件 UI 屏页面切换")
    add_body_paragraph(
        doc,
        "UI 屏页面切换分为默认切换逻辑和上位机配置切换逻辑。默认逻辑用于设备上电后的常规轮播；配置逻辑用于现场需要调整 Logo 页面、参数页面或总轮播时间时使用。",
        after=4,
        keep_next=True,
    )
    add_table(
        doc,
        ("切换逻辑", "执行规则"),
        [
            ("默认页面切换", "设备上电后，UI 屏首先显示 Logo 页面，默认显示 5 秒。Logo 页面结束后，系统自动进入参数页面轮播，设备第一页与设备第二页按默认时间依次切换，默认每页显示 5 秒。参数页面累计轮播达到 1 分钟后，系统再次切换至 Logo 页面显示 5 秒，并按上述逻辑循环执行。"),
            ("上位机配置切换", "页面切换时间可通过上位机配置。上位机下发一次页面时间配置命令后，UI 屏将按照配置的 Logo 显示时间、参数页面显示时间及总轮播时间执行页面切换。配置的总轮播时间结束后，系统自动恢复默认切换逻辑。"),
        ],
        [2300, 7060],
        font_size=8.9,
    )
    doc.add_page_break()


def insert_before_heading(doc, heading_text, builder):
    target = None
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == heading_text:
            target = paragraph
            break
    if target is None:
        raise RuntimeError(f"Target heading not found: {heading_text}")

    body = doc._body._element
    before = list(body)
    builder(doc)
    after = list(body)
    new_elements = [element for element in after if element not in before and element.tag != qn("w:sectPr")]

    for element in new_elements:
        body.remove(element)
    for element in new_elements:
        target._p.addprevious(element)


def main():
    for image in (
        IMG_OPERATION_LIVE,
        IMG_OPERATION_HISTORY,
        IMG_OPERATION_TIME_DIALOG,
        IMG_OPERATION_SEARCH_RESULT,
        IMG_SOFTWARE_UI,
        IMG_FIRMWARE_FLOW,
        IMG_FIRMWARE_LOGO,
        IMG_FIRMWARE_PAGE1,
        IMG_FIRMWARE_PAGE2,
    ):
        if not image.exists():
            raise FileNotFoundError(image)

    doc = Document(SOURCE)
    if any(paragraph.text.strip().startswith("4.1.9 软件 OHB 设备 UI") for paragraph in doc.paragraphs):
        raise RuntimeError("Document already contains section 4.1.9.")

    renumber_following_figures(doc)
    insert_before_heading(doc, "4.1.4 Set 详细信息", append_operation_log_detail_section)
    insert_before_heading(doc, "4.2 Cabinet 界面介绍", append_home_ohb_ui_section)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()

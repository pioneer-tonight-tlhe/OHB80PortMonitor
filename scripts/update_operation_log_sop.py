from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


SOURCE_DOCX = Path(r"C:\Users\Lenovo\Desktop\OHB80PortMonitor软件SOP_.docx")
OUTPUT_DOCX = Path(r"D:\Project\CYTC_Project\OHB80PortMonitor\OHB80PortMonitor_SOP_4_4_to_4_6_updated.docx")

IMAGES = {
    "live": Path(r"C:\Users\Lenovo\AppData\Local\Temp\codex-clipboard-df2c2bf7-bed9-4310-8ea9-079139f62c14.png"),
    "history": Path(r"C:\Users\Lenovo\AppData\Local\Temp\codex-clipboard-3565c85e-1e40-4335-a2d9-ce75d1f72bd6.png"),
    "time_popup": Path(r"C:\Users\Lenovo\AppData\Local\Temp\codex-clipboard-aed707b9-e068-444d-9a60-9f4fb491108c.png"),
    "keyword": Path(r"C:\Users\Lenovo\AppData\Local\Temp\codex-clipboard-cf0ba8f8-b860-4973-b59f-defa4eaded8b.png"),
    "comm_live": Path(r"C:\Users\Lenovo\AppData\Local\Temp\codex-clipboard-95bdae63-5417-4391-b692-1bbcae6f6c86.png"),
    "comm_history": Path(r"C:\Users\Lenovo\AppData\Local\Temp\codex-clipboard-58205673-e56b-44cf-a707-5d5e1306fb61.png"),
    "alarm_live": Path(r"C:\Users\Lenovo\AppData\Local\Temp\codex-clipboard-ed25b64f-e7fe-4116-9a6a-5c08d8f8d520.png"),
    "alarm_history": Path(r"C:\Users\Lenovo\AppData\Local\Temp\codex-clipboard-50b9c613-01bb-4311-991d-67d47cb76018.png"),
}
GENERATED_DIR = Path(r"D:\Project\CYTC_Project\OHB80PortMonitor\generated_doc_assets")
TIME_POPUP_NUMBERED = GENERATED_DIR / "operation_log_time_popup_numbered.png"

FONT = "Microsoft YaHei"
BODY_COLOR = RGBColor(0x21, 0x21, 0x21)
CAPTION_COLOR = RGBColor(0x5A, 0x5A, 0x5A)
HEADER_FILL = "D9EAF7"
HEADER_TEXT = RGBColor(0x1F, 0x4D, 0x78)
BORDER = "9CC2E5"


def set_run_font(run, size_pt=10.5, bold=False, color=BODY_COLOR):
    run.font.name = FONT
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), FONT)
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.font.color.rgb = color


def add_body(doc, text="", after_pt=6):
    para = doc.add_paragraph(style="Normal")
    para.paragraph_format.space_after = Pt(after_pt)
    para.paragraph_format.line_spacing = 1.25
    if text:
        run = para.add_run(text)
        set_run_font(run)
    return para


def add_heading(doc, text, style):
    para = doc.add_paragraph(style=style)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if not para.runs:
        run = para.add_run(text)
    else:
        para.runs[0].text = text
        run = para.runs[0]
    if style == "SOP Heading 2":
        set_run_font(run, size_pt=14, bold=True, color=HEADER_TEXT)
    else:
        set_run_font(run, size_pt=12, bold=True, color=HEADER_TEXT)
    return para


def add_caption(doc, text):
    para = doc.add_paragraph(style="Normal")
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(8)
    run = para.add_run(text)
    set_run_font(run, size_pt=9, color=CAPTION_COLOR)
    return para


def add_image(doc, path, width_in=6.45):
    para = doc.add_paragraph(style="Normal")
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run()
    run.add_picture(str(path), width=Inches(width_in))
    return para


def set_cell_text(cell, text, bold=False, fill=None, color=BODY_COLOR, size_pt=9.5):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.15
    para.clear()
    run = para.add_run(text)
    set_run_font(run, size_pt=size_pt, bold=bold, color=color)
    if fill:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = tc_pr.find(qn("w:shd"))
        if shd is None:
            shd = OxmlElement("w:shd")
            tc_pr.append(shd)
        shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), BORDER)


def set_table_width(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_grid = table._tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, width in enumerate(widths_dxa):
            set_cell_width(row.cells[idx], width)


def add_table(doc, headers, rows, widths_dxa, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths_dxa)
    set_table_borders(table)
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, fill=HEADER_FILL, color=HEADER_TEXT, size_pt=font_size)
    for values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            set_cell_text(cells[idx], value, size_pt=font_size)
    add_body(doc, "", after_pt=4)
    return table


def create_time_popup_numbered():
    GENERATED_DIR.mkdir(parents=True, exist_ok=True)
    image = Image.open(IMAGES["time_popup"]).convert("RGBA")
    draw = ImageDraw.Draw(image)
    try:
        font = ImageFont.truetype("arial.ttf", 18)
    except OSError:
        font = ImageFont.load_default()

    labels = [
        (1, 18, 58),
        (2, 18, 158),
        (3, 462, 396),
    ]
    radius = 16
    for number, x, y in labels:
        box = (x - radius, y - radius, x + radius, y + radius)
        draw.ellipse(box, fill=(237, 64, 64, 255))
        text = str(number)
        text_box = draw.textbbox((0, 0), text, font=font)
        tw = text_box[2] - text_box[0]
        th = text_box[3] - text_box[1]
        draw.text((x - tw / 2, y - th / 2 - 1), text, fill=(255, 255, 255, 255), font=font)

    image.convert("RGB").save(TIME_POPUP_NUMBERED)
    return TIME_POPUP_NUMBERED


def append_operation_log_section(doc):
    time_popup_image = create_time_popup_numbered()

    add_body(doc, "", after_pt=2)
    add_heading(doc, "4.4 运行日志界面介绍", "SOP Heading 2")
    add_body(
        doc,
        "运行日志界面用于查看软件运行过程中产生的实时日志和历史日志。该界面包含 Live Log 与 History 两个子标签：Live Log 用于查看当前实时产生的记录，History 用于按日志类型、时间范围和关键字查询历史记录。",
    )

    add_heading(doc, "4.4.1 Live Log 实时界面字段说明", "SOP Heading 3")
    add_body(doc, "Live Log 实时界面如下：", after_pt=4)
    add_image(doc, IMAGES["live"], width_in=6.45)
    add_caption(doc, "图 4-27 Live Log 实时界面区域与编号标注")
    add_table(
        doc,
        ["编号", "配置项", "功能说明"],
        [
            ("1", "运行日志导航栏", "包含 Live Log 和 History 两个标签。Live Log 用于查看实时产生的运行日志；History 用于进入历史日志查询界面。"),
            ("2", "日志表头", "实时日志列表的字段表头，包括 Occur Time、Log Type、Description。Occur Time 表示日志发生时间；Log Type 表示日志类型，如 Message 或 Error；Description 表示日志详细内容。"),
        ],
        [700, 2400, 4900],
    )

    add_heading(doc, "4.4.2 History 历史界面字段说明", "SOP Heading 3")
    add_body(doc, "History 历史界面用于按条件查询历史运行日志，界面如下：", after_pt=4)
    add_image(doc, IMAGES["history"], width_in=6.45)
    add_caption(doc, "图 4-28 History 历史界面区域与编号标注")
    add_table(
        doc,
        ["编号", "配置项", "功能说明"],
        [
            ("1", "所有条件启用", "All 勾选框用于启用所有历史日志查询条件。勾选后，日志类型条件和时间范围条件可一并参与查询。"),
            ("2", "日志类型条件", "Log Type 用于按日志类型筛选历史记录，可选择 Message、Error 等类型。勾选该条件后，系统只查询所选类型的日志。"),
            ("3", "时间范围条件", "Record Time 用于按日志发生时间筛选历史记录。勾选该条件后，可通过 Set 按钮设置开始时间和结束时间。"),
            ("4", "关键字查询条件", "关键字输入框用于输入需要查询的内容，例如设备编号、二维码、报警码或描述中的关键字；Search、Pre、Next、0/0、Record No.、Jump 用于执行搜索和定位记录。"),
            ("5", "分页按钮", "用于在历史日志查询结果中翻页查看记录，包括上一页、下一页、页码、Page 当前页/总页数、跳页输入框和 GO 按钮。"),
        ],
        [700, 2400, 4900],
    )

    add_heading(doc, "4.4.3 时间范围设置弹框", "SOP Heading 3")
    add_body(doc, "点击 History 界面中的 Set 按钮后，会弹出 Set Record Time 时间设置弹框，如下：", after_pt=4)
    add_image(doc, time_popup_image, width_in=5.05)
    add_caption(doc, "图 4-29 Set Record Time 时间设置弹框区域与编号标注")
    add_table(
        doc,
        ["编号", "配置项", "功能说明"],
        [
            ("1", "Enable Start Time", "启用开始时间。勾选后，下方日期和时间作为历史日志查询范围的起始时间。"),
            ("2", "Enable End Time", "启用结束时间。勾选后，下方日期和时间作为历史日志查询范围的结束时间。"),
            ("3", "OK / Cancel", "OK 用于确认当前时间范围设置并返回 History 界面；Cancel 用于取消本次时间设置，不保存当前修改。"),
        ],
        [700, 2400, 4900],
    )
    add_body(
        doc,
        "使用时间范围查询时，应先在 History 界面勾选 Record Time 条件，再通过 Set 设置开始时间和结束时间。只启用开始时间时，查询起始时间之后的记录；只启用结束时间时，查询结束时间之前的记录；同时启用开始时间和结束时间时，查询两者之间的记录。",
    )

    add_heading(doc, "4.4.4 关键字搜索与记录定位", "SOP Heading 3")
    add_body(
        doc,
        "在 History 界面的关键字输入框中输入需要查询的关键字，点击 Search 按钮即可搜索到对应记录。下图为搜索关键字 12001 后的显示效果：",
        after_pt=4,
    )
    add_image(doc, IMAGES["keyword"], width_in=6.45)
    add_caption(doc, "图 4-30 搜索关键字 12001 后的历史日志区域与编号标注")
    add_body(
        doc,
        "搜索后，系统会在当前查询结果中定位包含关键字的记录。图中黄色高亮区域表示当前被选中的第一条匹配记录，绿色高亮区域表示本次查询命中的其他记录。",
    )
    add_table(
        doc,
        ["编号", "配置项", "功能说明"],
        [
            ("1", "条件查询被选中的第一条记录", "输入关键字 12001 并点击 Search 后，系统会定位到当前命中的第一条记录，并以黄色高亮显示。"),
            ("2", "添加查询到的其他记录", "除当前选中的第一条记录外，列表中其他包含关键字 12001 的记录会以绿色高亮显示，便于继续查看。"),
        ],
        [700, 2400, 4900],
    )
    add_body(doc, "关键字搜索控件说明如下：", after_pt=4)
    add_table(
        doc,
        ["编号", "配置项", "功能说明"],
        [
            ("1", "Search", "在输入框中输入需要查询的关键字后，点击 Search 按钮即可搜索到对应记录。"),
            ("2", "Pre", "切换到上一条匹配记录，用于向前查看本次搜索命中的日志。"),
            ("3", "Next", "切换到下一条匹配记录，用于向后查看本次搜索命中的日志。"),
            ("4", "0/0", "显示当前匹配记录序号 / 匹配记录总数。未搜索到记录时显示 0/0；搜索到记录后会显示类似 1/60 的结果。"),
            ("5", "Record No.", "记录序号输入框，用于输入需要跳转的匹配记录编号。"),
            ("6", "Jump", "根据 Record No. 输入的编号，直接跳转到对应的匹配记录。"),
        ],
        [700, 2400, 4900],
    )

    add_heading(doc, "4.4.5 分页按钮说明", "SOP Heading 3")
    add_body(doc, "History 界面底部的分页按钮用于在历史日志查询结果中按页切换，适合日志数量较多时使用。")
    add_table(
        doc,
        ["编号", "配置项", "功能说明"],
        [
            ("1", "<", "切换到上一页历史日志记录。"),
            ("2", "页码按钮", "直接切换到指定页。当前页按钮会以高亮状态显示。"),
            ("3", ">", "切换到下一页历史日志记录。"),
            ("4", "Page1/6", "显示当前页码和总页数，格式为 当前页 / 总页数。"),
            ("5", "跳至 1 页", "输入需要跳转的页码。"),
            ("6", "GO", "执行跳页操作，跳转到输入的页码。"),
        ],
        [700, 2400, 4900],
    )


def append_comm_log_section(doc):
    add_body(doc, "", after_pt=2)
    add_heading(doc, "4.5 Comm. 通讯日志界面", "SOP Heading 2")
    add_body(
        doc,
        "概述：Comm. 通讯日志界面用于实时监控上位机软件对设备下发的通讯指令，并记录指令下发时间、指令名称、执行状态、重发次数和详细通讯内容。该界面包含 Live Log 和 History 两个子标签，Live Log 用于查看实时通讯日志，History 用于按条件查询历史通讯日志。",
    )

    add_heading(doc, "4.5.1 Live Log 实时通讯日志界面", "SOP Heading 3")
    add_body(doc, "Comm. 的 Live Log 实时通讯日志界面如下：", after_pt=4)
    add_image(doc, IMAGES["comm_live"], width_in=6.45)
    add_caption(doc, "图 4-31 Comm. 实时通讯日志界面区域与编号标注")
    add_table(
        doc,
        ["编号", "配置项", "功能说明"],
        [
            ("1", "导航栏", "包含 Live Log 和 History 两个标签。Live Log 用于查看实时通讯日志；History 用于进入历史通讯日志查询界面。"),
            ("2", "日志表头", "通讯日志列表的字段表头，包括 QRCode、Send Time、Command ID、Description。QRCode 表示设备二维码编号；Send Time 表示指令下发时间；Command ID 表示通讯指令名称；Description 表示指令参数或通讯返回内容。"),
            ("3", "设备 QRCode", "显示接收通讯指令的设备 QRCode。该列按照 QRCode 的顺序排列，便于按设备编号快速查看通讯状态。"),
        ],
        [700, 2400, 4900],
    )

    add_heading(doc, "4.5.2 History 历史通讯日志界面", "SOP Heading 3")
    add_body(doc, "Comm. 的 History 历史通讯日志界面用于按条件查询已记录的通讯指令，如下：", after_pt=4)
    add_image(doc, IMAGES["comm_history"], width_in=6.45)
    add_caption(doc, "图 4-32 Comm. 历史通讯日志界面区域与编号标注")
    add_table(
        doc,
        ["编号", "配置项", "功能说明"],
        [
            ("1", "开启全部查询条件", "All 勾选框用于启用全部历史通讯日志查询条件。勾选后，QRCode、Command ID、Exec Status、Retry Count 和 Send Time 等条件可一并参与查询。"),
            ("2", "根据 QRCode 查询", "QRCode 条件用于按设备二维码编号查询通讯日志。选择或输入设备 QRCode 后，系统只显示对应设备的通讯记录。"),
            ("3", "根据指令名称查询", "Command ID 条件用于按通讯指令名称查询记录，例如 ReadBoardEnable 等指令。"),
            ("4", "根据指令执行状态查询", "Exec Status 条件用于按指令执行状态查询记录，例如 Success、Failed 等状态。"),
            ("5", "根据指令重发次数查询", "Retry Count 条件用于按指令重发次数查询记录，可用于定位重复发送或通讯不稳定的指令。"),
            ("6", "根据时间区间查询", "Send Time 条件用于按通讯指令下发时间范围查询记录。点击 Set 按钮可设置开始时间和结束时间。"),
            ("7", "分页按钮", "用于在历史通讯日志查询结果中翻页查看记录，包括上一页、下一页、页码、Page 当前页/总页数、跳页输入框和 GO 按钮。"),
        ],
        [700, 2400, 4900],
    )


def append_alarm_log_section(doc):
    add_body(doc, "", after_pt=2)
    add_heading(doc, "4.6 Alarm 警报日志界面", "SOP Heading 2")
    add_body(
        doc,
        "概述：Alarm 警报日志界面用于实时监控设备出现的警报，并记录警报级别、发生时间、设备 QRCode、警报类型、是否解决、解决时间和警报描述。该界面包含 Live Log 和 History 两个子标签，Live Log 用于查看实时警报日志，History 用于按条件查询历史警报日志。",
    )

    add_heading(doc, "4.6.1 Live Log 实时警报日志界面", "SOP Heading 3")
    add_body(doc, "Alarm 的 Live Log 实时警报日志界面如下：", after_pt=4)
    add_image(doc, IMAGES["alarm_live"], width_in=6.45)
    add_caption(doc, "图 4-33 Alarm 实时警报日志界面区域与编号标注")
    add_table(
        doc,
        ["编号", "配置项", "功能说明"],
        [
            ("1", "导航栏", "包含 Live Log 和 History 两个标签。Live Log 用于查看实时警报日志；History 用于进入历史警报日志查询界面。"),
            ("2", "日志表头", "警报日志列表的字段表头，包括 Alarm Level、Occur Time、QRCode、Alarm Type、Is Resolved、Resolve Time、Description。Alarm Level 表示警报级别；Occur Time 表示警报发生时间；Alarm Type 表示警报类型；Is Resolved 表示警报是否解决；Resolve Time 表示警报解决时间；Description 表示警报详细描述。"),
            ("3", "设备 QRCode", "显示发生警报的设备 QRCode。该列按照 QRCode 的顺序排列，便于按设备编号快速查看设备警报。"),
        ],
        [700, 2400, 4900],
    )

    add_heading(doc, "4.6.2 History 历史警报日志界面", "SOP Heading 3")
    add_body(doc, "Alarm 的 History 历史警报日志界面用于按条件查询已记录的警报，如下：", after_pt=4)
    add_image(doc, IMAGES["alarm_history"], width_in=6.45)
    add_caption(doc, "图 4-34 Alarm 历史警报日志界面区域与编号标注")
    add_table(
        doc,
        ["编号", "配置项", "功能说明"],
        [
            ("1", "开启所有查询条件", "All 勾选框用于启用全部历史警报日志查询条件。勾选后，QRCode、Alarm Level、Alarm Type、Is Resolved、Start Time 和 Resolved Time 等条件可一并参与查询。"),
            ("2", "根据设备 QRCode 查询", "QRCode 条件用于按设备二维码编号查询警报记录。选择或输入设备 QRCode 后，系统只显示对应设备的警报日志。"),
            ("3", "根据警报级别查询", "Alarm Level 条件用于按警报级别查询记录，例如 warn、Error、Fatal。"),
            ("4", "根据警报类型查询", "Alarm Type 条件用于按警报类型查询记录，例如 Device Offline、SH85 Abnormal 等。"),
            ("5", "根据警报是否解决查询", "Is Resolved 条件用于按警报解决状态查询记录，可选择已解决或未解决状态。"),
            ("6", "根据警报发生/解决时间区间搜索条件", "Start Time 用于按警报发生时间范围查询；Resolved Time 用于按警报解决时间范围查询。点击对应 Set 按钮可设置开始时间和结束时间。"),
            ("7", "分页按钮", "用于在历史警报日志查询结果中翻页查看记录，包括上一页、下一页、页码、Page 当前页/总页数、跳页输入框和 GO 按钮。"),
        ],
        [700, 2400, 4900],
    )

    add_heading(doc, "4.6.3 警报级别说明", "SOP Heading 3")
    add_table(
        doc,
        ["警报级别", "功能说明"],
        [
            ("warn", "普通报警，不影响设备正常运行。"),
            ("Error", "错误，可能影响设备正常运行。"),
            ("Fatal", "非常严重的错误。"),
        ],
        [1800, 6200],
    )

    add_heading(doc, "4.6.4 警报类型说明", "SOP Heading 3")
    add_table(
        doc,
        ["枚举值", "编号", "显示名称", "告警等级", "含义"],
        [
            ("DeviceOffline", "2001", "Device Offline", "Error", "设备离线"),
            ("TemperatureSensorAbnormal", "3001", "Temperature Sensor Abnormal", "Error", "温度传感器异常"),
            ("HumiditySensorAbnormal", "3002", "Humidity Sensor Abnormal", "Error", "湿度传感器异常"),
            ("SH85SelfCheckActionFailed", "4100", "SH85 Self Check Action Failed", "Error", "SH85 自检动作失败"),
            ("SH85AcceptanceHumidityExceeded", "4106", "SH85 Acceptance Humidity Exceeded", "Fatal", "SH85 验收湿度超限"),
            ("SH85AcceptanceSensorCommError", "4107", "SH85 Acceptance Sensor Comm Error", "Fatal", "SH85 验收传感器通信错误"),
            ("SH85AcceptanceThresholdParamError", "4108", "SH85 Acceptance Threshold Param Error", "Fatal", "SH85 验收阈值参数错误"),
            ("SH85Abnormal", "5003", "SH85 Abnormal", "Error", "SH85 异常"),
            ("HumidityNotReached", "5101", "Humidity Not Reached", "Error", "湿度未达标"),
        ],
        [2400, 800, 2800, 1000, 2000],
        font_size=8.5,
    )


def main():
    for image in IMAGES.values():
        if not image.exists():
            raise FileNotFoundError(image)
    doc = Document(str(SOURCE_DOCX))
    append_operation_log_section(doc)
    append_comm_log_section(doc)
    append_alarm_log_section(doc)
    doc.save(str(OUTPUT_DOCX))
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from pathlib import Path
from datetime import datetime

root = Path(r"D:\Project\CYTC_Project\OHB80PortMonitor\OHB80PortMonitor_V_1_0_0")
assets = root / "docs" / "manual_assets"
out_doc = root / "docs" / "OHB80PortMonitor_用户使用手册_v1.1_优化版.docx"

img = {
    "main_a": assets / "img_01.png",
    "guest": assets / "img_02.png",
    "login": assets / "img_03.png",
    "user": assets / "img_04.png",
    "main_b": assets / "img_05.png",
    "set_view": assets / "img_06.png",
    "foup_view": assets / "img_07.png",
    "config": assets / "img_08.png",
    "sh85": assets / "img_09.png",
}
for p in img.values():
    if not p.exists():
        raise FileNotFoundError(str(p))

doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(1)
sec.bottom_margin = Inches(1)
sec.left_margin = Inches(1)
sec.right_margin = Inches(1)

def set_run_font(run, size=11, bold=False, color=None):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

def add_para(text="", size=11, bold=False, align=None, color=None, spacing_after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_after = Pt(spacing_after)
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, color=color)
    return p

def add_heading(text, level=1):
    h = doc.add_heading(level=level)
    r = h.add_run(text)
    set_run_font(r, size=16 if level==1 else 13, bold=True, color=RGBColor(31,73,125))
    h.paragraph_format.space_before = Pt(8)
    h.paragraph_format.space_after = Pt(6)

fig_no = 1
def add_figure(path, caption, width=6.2):
    global fig_no
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.paragraph_format.space_after = Pt(3)
    p.add_run().add_picture(str(path), width=Inches(width))

    cp = doc.add_paragraph()
    cp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    cp.paragraph_format.space_after = Pt(8)
    r = cp.add_run(f"图{fig_no}  {caption}")
    set_run_font(r, size=10, bold=False, color=RGBColor(90,90,90))
    fig_no += 1

# Cover
add_para("OHB80PortMonitor 用户使用手册", size=24, bold=True, align=WD_PARAGRAPH_ALIGNMENT.CENTER, color=RGBColor(31,73,125), spacing_after=10)
add_para("版本：V1.1（优化版）", size=13, align=WD_PARAGRAPH_ALIGNMENT.CENTER, spacing_after=2)
add_para(f"日期：{datetime.now().strftime('%Y-%m-%d')}", size=12, align=WD_PARAGRAPH_ALIGNMENT.CENTER, spacing_after=20)
add_para("适用对象：操作员 / 现场维护 / 培训人员", size=11, align=WD_PARAGRAPH_ALIGNMENT.CENTER, spacing_after=2)
add_para("适用软件：OHB80PortMonitor V1.0.0", size=11, align=WD_PARAGRAPH_ALIGNMENT.CENTER, spacing_after=20)

add_para("文档说明", size=14, bold=True, color=RGBColor(31,73,125), spacing_after=6)
add_para("本手册用于指导用户完成系统登录、首页监控查看、视图切换、配置页面基础操作及 SH85 自检功能使用。")
add_para("当前文档已使用你提供的 1~9 号截图进行配图。通信页、报警页、调试页等可在后续补图后继续扩展。", spacing_after=12)

# Revision table
tbl = doc.add_table(rows=1, cols=4)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = tbl.rows[0].cells
for i, t in enumerate(["版本", "日期", "变更内容", "作者"]):
    pr = hdr[i].paragraphs[0]
    pr.clear()
    rr = pr.add_run(t)
    set_run_font(rr, size=10, bold=True)
row = tbl.add_row().cells
vals = ["V1.0", "2026-05-28", "初版", "Codex"]
for i,v in enumerate(vals):
    row[i].paragraphs[0].clear()
    rr = row[i].paragraphs[0].add_run(v)
    set_run_font(rr, size=10)
row2 = tbl.add_row().cells
vals2 = ["V1.1", datetime.now().strftime('%Y-%m-%d'), "结构优化、步骤细化、排版统一", "Codex"]
for i,v in enumerate(vals2):
    row2[i].paragraphs[0].clear()
    rr = row2[i].paragraphs[0].add_run(v)
    set_run_font(rr, size=10)

# TOC-like list
doc.add_page_break()
add_heading("目录", 1)
for line in [
    "1. 软件界面总览",
    "2. 登录与权限管理",
    "3. 首页监控与视图切换",
    "4. 配置页面操作说明",
    "5. SH85 自检功能",
    "6. 使用建议与注意事项",
    "7. 常见问题与排查",
]:
    add_para(line, size=11, spacing_after=2)

doc.add_page_break()

# 1
add_heading("1. 软件界面总览", 1)
add_para("系统主界面由顶部状态提示、左侧导航菜单、中部参数表、底部轨道图四部分组成。")
add_figure(img["main_a"], "系统主界面总览（Home + Alarm）", width=6.2)

# 2
add_heading("2. 登录与权限管理", 1)
add_para("系统采用分级权限控制，建议每位用户使用自己的账号登录，便于审计与追踪。")
add_heading("2.1 未登录状态", 2)
add_para("未登录时为 Guest 级别，仅可进行基础查看。")
add_figure(img["guest"], "账号菜单（Guest）", width=2.8)

add_heading("2.2 登录操作", 2)
add_para("步骤 1：点击右上角用户图标，选择 Login。", spacing_after=2)
add_para("步骤 2：输入 Username 与 Password。", spacing_after=2)
add_para("步骤 3：点击 Login，认证通过后进入对应权限。", spacing_after=6)
add_figure(img["login"], "登录窗口", width=4.8)

add_heading("2.3 登录后账号菜单", 2)
add_para("登录后可执行 Change Password、Login New Account、Logout 等操作。")
add_figure(img["user"], "账号菜单（Normal）", width=2.8)

# 3
add_heading("3. 首页监控与视图切换", 1)
add_para("首页用于实时查看设备关键参数，建议作为日常巡检主页面。")
add_figure(img["main_b"], "首页监控总览", width=6.2)

add_para("关键区域说明：", bold=True, spacing_after=2)
for t in [
    "1) 顶部告警条：显示实时告警摘要。",
    "2) 左侧导航栏：Home / Config / Comm. / Alarm 快速切换。",
    "3) 参数表格：显示 QRCode、压力、流量、湿度、温度。",
    "4) 底部轨道图：设备位置及状态可视化。",
]:
    add_para(t, spacing_after=1)

add_heading("3.1 视图切换（Set View / Foup View）", 2)
add_para("通过右上角按钮可在 Set 视图与 Foup 视图间切换，便于不同场景观察。")
add_figure(img["set_view"], "Set View 示例", width=6.2)
add_figure(img["foup_view"], "Foup View 示例", width=6.2)

# 4
add_heading("4. 配置页面操作说明", 1)
add_para("配置页面按功能分区，可在顶部标签页快速定位各设置模块。")
add_figure(img["config"], "Config 页面总览", width=6.2)

add_para("常见模块：", bold=True, spacing_after=2)
for t in [
    "1) Idle Purge：空闲吹扫使能、时长、间隔。",
    "2) Pneumatic Valve：目标设备与气动阀压力设置。",
    "3) SH85 Periodic：定期自检开关、周期、状态与报告。",
    "4) SH85 Self-check：目标设备 ID 与手动自检触发。",
]:
    add_para(t, spacing_after=1)

# 5
add_heading("5. SH85 自检功能", 1)
add_para("SH85 自检分为“定期自检”和“手动自检”，建议先启用定期自检，再按需执行手动自检。")
add_figure(img["sh85"], "SH85 自检区域（Periodic + Manual）", width=6.0)
add_para("定期自检操作：", bold=True, spacing_after=2)
add_para("1) Enable Periodic Self-check 设为 true。", spacing_after=1)
add_para("2) 设置 Self-check Period（例如 5 min），点击 Set。", spacing_after=1)
add_para("3) 观察 Self-check Status 倒计时是否正常变化。", spacing_after=6)

add_para("手动自检操作：", bold=True, spacing_after=2)
add_para("1) 输入 Target Device ID。", spacing_after=1)
add_para("2) 点击 Check 触发自检（约 70 秒）。", spacing_after=1)
add_para("3) 结合顶部告警信息与状态回读判断结果。", spacing_after=6)

# 6
add_heading("6. 使用建议与注意事项", 1)
for t in [
    "1) 修改参数前，先核对目标设备 ID / QRCode，避免误下发。",
    "2) 大批量操作建议先在单设备验证，再执行全量设置。",
    "3) 观察顶部告警条与回读值，确认设置已生效。",
    "4) 现场网络波动时优先确认链路与设备在线状态。",
    "5) 建议定期更换账号密码并按权限最小化原则分配账号。",
]:
    add_para(t, spacing_after=2)

# 7
add_heading("7. 常见问题与排查", 1)
faq = [
    ("Q1：无法登录？", "检查用户名/密码是否正确，确认账号未被权限策略限制。"),
    ("Q2：设置后无变化？", "确认目标设备 ID 是否正确，查看网络状态与是否有告警拦截。"),
    ("Q3：自检无结果？", "确认设备在线，等待完整自检时长后再查看状态与告警信息。"),
    ("Q4：页面数据显示异常？", "优先刷新页面并核对通信链路，必要时重新登录。"),
]
for q,a in faq:
    add_para(q, bold=True, spacing_after=1)
    add_para(a, spacing_after=4)

add_para("—— 结束 ——", align=WD_PARAGRAPH_ALIGNMENT.CENTER, color=RGBColor(100,100,100), spacing_after=0)

doc.save(out_doc)
print(out_doc)
